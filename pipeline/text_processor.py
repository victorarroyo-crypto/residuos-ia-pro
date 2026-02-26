"""
PROCESADOR DE DOCUMENTOS DE TEXTO - ResidusIA Pro
===================================================
Procesa documentos de texto plano y enriquecido:
  - DOCX (Microsoft Word)
  - TXT (texto plano)
  - HTML (páginas web guardadas)
  - MD (Markdown)

Extrae el texto, lo clasifica y lo fragmenta con el mismo
chunker semántico que los PDFs.
"""

import hashlib
import io
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from .pdf_pipeline import (
    DocType,
    DocumentChunk,
    PageContent,
    PDFNature,
    PipelineConfig,
    ProcessedDocument,
)
from .classifier_chunker import DocumentClassifier, SemanticChunker
from .config import EmbeddingService
from .storage import StorageService

logger = logging.getLogger(__name__)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extrae texto de un archivo DOCX."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error extrayendo texto de DOCX: {e}")
        return ""


def extract_html_text(file_bytes: bytes) -> str:
    """Extrae texto de un archivo HTML eliminando tags."""
    text = file_bytes.decode("utf-8", errors="replace")
    # Eliminar scripts y styles
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Convertir <br>, <p>, <div>, <li> a saltos de línea
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", text, flags=re.IGNORECASE)
    # Eliminar todos los tags restantes
    text = re.sub(r"<[^>]+>", "", text)
    # Limpiar entidades HTML comunes
    text = text.replace("&nbsp;", " ").replace("&amp;", "&")
    text = text.replace("&lt;", "<").replace("&gt;", ">")
    text = text.replace("&quot;", '"')
    # Limpiar espacios excesivos
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_txt_text(file_bytes: bytes) -> str:
    """Extrae texto de un archivo de texto plano."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


# Mapa de extractores por extensión
TEXT_EXTRACTORS = {
    "docx": extract_docx_text,
    "doc": extract_docx_text,
    "txt": extract_txt_text,
    "html": extract_html_text,
    "htm": extract_html_text,
    "md": extract_txt_text,
}


class TextProcessor:
    """
    Procesa documentos de texto (DOCX, TXT, HTML, MD).
    Usa el mismo flujo que PDFs: clasificar → chunking → embeddings → storage.
    """

    def __init__(self, config: PipelineConfig):
        self.config = config
        self.classifier = DocumentClassifier(config)
        self.chunker = SemanticChunker(config)
        self.embedder = EmbeddingService(config)
        self.storage = StorageService(config)

    async def process(
        self,
        file_bytes: bytes,
        filename: str,
        client_id: str,
        project_id: Optional[str] = None,
        rag_scope: str = "general",
    ) -> ProcessedDocument:
        """Procesa un documento de texto y lo indexa."""
        ext = Path(filename).suffix.lower().lstrip(".")
        extractor = TEXT_EXTRACTORS.get(ext)

        if not extractor:
            raise ValueError(f"Formato de texto no soportado: .{ext}")

        doc_id = self._generate_doc_id(file_bytes, client_id)

        # Extraer texto
        text = extractor(file_bytes)
        if not text.strip():
            raise ValueError(f"No se pudo extraer texto de {filename}")

        logger.info(f"[{filename}] Texto extraído: {len(text)} caracteres")

        # Crear PageContent simulado (el chunker espera esta estructura)
        pages = [
            PageContent(
                page_num=1,
                text=text,
                tables=[],
                images=[],
                nature=PDFNature.DIGITAL,
                confidence=1.0,
            )
        ]

        # Clasificar
        doc_type = await self.classifier.classify(pages, filename, project_id=project_id)
        logger.info(f"[{filename}] Tipo detectado: {doc_type}")

        # Chunking semántico
        chunks = await self.chunker.chunk(pages, doc_type, doc_id, filename=filename)
        logger.info(f"[{filename}] {len(chunks)} chunks generados")

        # Embeddings
        chunks = await self.embedder.embed_all(chunks)

        # Subir archivo original a Storage
        storage_path = await self.storage.upload_file(
            file_bytes, filename, client_id, doc_type
        )

        # Construir documento procesado
        processed = ProcessedDocument(
            doc_id=doc_id,
            client_id=client_id,
            original_filename=filename,
            doc_type=doc_type,
            nature=PDFNature.DIGITAL,
            total_pages=1,
            chunks=chunks,
            tables_found=0,
            was_encrypted=False,
            ocr_applied=False,
            ocr_avg_confidence=1.0,
            extraction_warnings=[],
            metadata={"rag_scope": rag_scope, "format": ext},
            storage_path=storage_path,
        )

        # Guardar en Supabase
        is_knowledge = self.storage._is_knowledge(processed)
        processed.supabase_doc_id = await self.storage.save_to_supabase(processed)
        await self.storage.save_chunks_to_supabase(
            chunks, processed.supabase_doc_id,
            is_knowledge=is_knowledge,
            project_id=None if is_knowledge else client_id,
        )

        logger.info(f"[{filename}] Procesado y almacenado. Doc ID: {doc_id}")
        return processed

    def _generate_doc_id(self, file_bytes: bytes, client_id: str) -> str:
        h = hashlib.sha256(file_bytes + client_id.encode()).hexdigest()[:16]
        return f"doc_{h}"
