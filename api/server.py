"""
ResidusIA Pro - API Server
Expone el pipeline de procesamiento de documentos via HTTP.
"""

import gc
import ipaddress
import json
import os
import socket
import sys
import time
from contextlib import asynccontextmanager
from typing import Optional

import asyncio
import logging
from urllib.parse import urlparse

import httpx
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field, field_validator
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

_UUID_PATTERN = r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$"
_VALID_TIERS = {"standard", "pro_plus"}
_VALID_MODELS = {
    "claude-opus-4-6", "claude-sonnet-4", "claude-haiku-4-5",
    "gpt-5.2", "gpt-5", "o3", "o4-mini", "gpt-5-mini",
    "gemini-2.5-pro", "gemini-2.5-flash",
}
_VALID_AGENTS = {"aai", "contratos", "facturas", "registro", "normativo"}

logging.basicConfig(level=logging.INFO, stream=sys.stdout)
logger = logging.getLogger("residusia")
logger.setLevel(logging.INFO)

# Silence noisy HTTP client loggers — these flood Railway logs as false "errors"
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)

# Ensure the project root is in the Python path (works locally and in Docker)
_project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

from pipeline import UnifiedIngestionService, PipelineConfigImpl, RAGScopingService, RAGScope
from pipeline.cost_guard import CostGuard, calculate_cost, get_provider, MODEL_PRICING
from pipeline.model_router import ModelRouter, MODEL_API_IDS, MODEL_PROVIDERS, SERVICE_DEFAULTS


service: UnifiedIngestionService | None = None
rag_service: RAGScopingService | None = None
_config: PipelineConfigImpl | None = None
_cost_guard: CostGuard | None = None
_model_router: ModelRouter | None = None

# Strong references to background tasks so GC doesn't kill them.
# See: https://docs.python.org/3/library/asyncio-task.html#asyncio.create_task
_background_tasks: set[asyncio.Task] = set()


@asynccontextmanager
async def lifespan(app: FastAPI):
    global service, rag_service, _config, _cost_guard, _model_router

    supabase_url = os.environ.get(
        "SUPABASE_URL", os.environ.get("NEXT_PUBLIC_SUPABASE_URL", "")
    )
    supabase_key = os.environ.get(
        "SUPABASE_SERVICE_KEY", os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    )

    if not supabase_url:
        logger.error(
            "SUPABASE_URL / NEXT_PUBLIC_SUPABASE_URL no configurado. "
            "Los datos NO llegarán a Supabase."
        )
        raise RuntimeError("SUPABASE_URL no configurado")
    if not supabase_key:
        logger.error(
            "SUPABASE_SERVICE_KEY / SUPABASE_SERVICE_ROLE_KEY no configurado. "
            "Los datos NO llegarán a Supabase."
        )
        raise RuntimeError("SUPABASE_SERVICE_ROLE_KEY no configurado")

    logger.info(f"Supabase URL: {supabase_url[:40]}...")
    logger.info("Supabase service key: configurada ✓")

    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        logger.warning("GEMINI_API_KEY no configurada — fallback a Gemini deshabilitado")

    _config = PipelineConfigImpl(
        anthropic_api_key=os.environ["ANTHROPIC_API_KEY"],
        openai_api_key=os.environ["OPENAI_API_KEY"],
        gemini_api_key=gemini_key,
        supabase_url=supabase_url,
        supabase_service_key=supabase_key,
    )
    service = UnifiedIngestionService(_config)
    rag_service = RAGScopingService(_config)

    # Cost Guard & Model Router
    _cost_guard = CostGuard(supabase_url, supabase_key)
    _model_router = ModelRouter(
        anthropic_api_key=_config.anthropic_api_key,
        openai_api_key=_config.openai_api_key,
        gemini_api_key=gemini_key,
        cost_guard=_cost_guard,
    )

    # Share ModelRouter with agent LLM layer
    from pipeline.agents.llm import init_model_router
    init_model_router(_model_router, _cost_guard)
    logger.info("CostGuard + ModelRouter inicializados ✓")

    # Clean up zombie syncs left by previous server instances (e.g. redeploy)
    try:
        from datetime import datetime as _dt, timezone as _tz
        from supabase._async.client import create_client as acreate_client
        _startup_sb = await acreate_client(supabase_url, supabase_key)
        _zombie_result = await (
            _startup_sb.table("gdrive_sync_log")
            .update({
                "status": "error",
                "completed_at": _dt.now(_tz.utc).isoformat(),
                "error_message": "Sync interrumpido por reinicio del servidor.",
            })
            .eq("status", "running")
            .execute()
        )
        _zombie_count = len(_zombie_result.data) if _zombie_result.data else 0
        if _zombie_count:
            logger.info("Startup: marked %d zombie sync(s) as error", _zombie_count)
    except Exception as e:
        logger.warning("Startup: could not clean zombie syncs: %s", e)

    yield


app = FastAPI(
    title="ResidusIA Pro API",
    description="Pipeline de procesamiento de documentos de residuos industriales",
    version="1.0.0",
    lifespan=lifespan,
)

# ── Rate Limiting ────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter


@app.exception_handler(RateLimitExceeded)
async def _rate_limit_handler(request: Request, exc: RateLimitExceeded):
    return JSONResponse(
        status_code=429,
        content={
            "detail": "Demasiadas solicitudes. Intenta de nuevo en unos momentos.",
            "retry_after": exc.detail,
        },
    )


_frontend_url = os.environ.get("FRONTEND_URL", "")
_cors_origins = [_frontend_url] if _frontend_url else []
if os.environ.get("ENVIRONMENT", "development") != "production":
    _cors_origins.append("http://localhost:3000")

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-API-Key"],
)

# ── API Key authentication ───────────────────────────────────────
# When PIPELINE_API_KEY is set, all endpoints (except /health) require
# the header X-API-Key to match.  If PIPELINE_API_KEY is empty/unset,
# authentication is skipped (local development).
_PIPELINE_API_KEY = os.environ.get("PIPELINE_API_KEY", "")
if _PIPELINE_API_KEY:
    logger.info("PIPELINE_API_KEY configured — API key authentication enabled ✓")
else:
    logger.warning("PIPELINE_API_KEY not set — API endpoints are UNPROTECTED")


@app.middleware("http")
async def _verify_api_key(request: Request, call_next):
    if request.url.path == "/health" or request.method == "OPTIONS":
        return await call_next(request)
    if _PIPELINE_API_KEY:
        provided = request.headers.get("x-api-key", "")
        if provided != _PIPELINE_API_KEY:
            return JSONResponse(
                status_code=401,
                content={"detail": "Invalid or missing API key"},
            )
    return await call_next(request)


@app.get("/health")
async def health():
    return {"status": "ok", "service": "residusia-pro-api"}


@app.post("/api/ingest")
@limiter.limit("10/minute")
async def ingest_document(
    request: Request,
    file: UploadFile = File(default=None),
    file_url: str = Form(default=None),
    storage_path: str = Form(default=None),
    filename: str = Form(default=None),
    project_id: str = Form(default=None),
    rag_scope: str = Form(default=None),
    password: str = Form(default=None),
):
    if service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    if not file and not file_url and not storage_path:
        raise HTTPException(
            status_code=400,
            detail="Debes enviar un archivo (`file`), una URL (`file_url`) o un `storage_path`.",
        )

    def _check_ssrf(url: str) -> None:
        """Block requests to private/internal IP addresses."""
        parsed = urlparse(url)
        hostname = parsed.hostname
        if not hostname:
            raise HTTPException(status_code=400, detail="URL sin hostname valido")
        try:
            resolved = socket.getaddrinfo(hostname, None, socket.AF_UNSPEC, socket.SOCK_STREAM)
            for family, _, _, _, sockaddr in resolved:
                ip = ipaddress.ip_address(sockaddr[0])
                if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                    raise HTTPException(
                        status_code=400,
                        detail="No se permiten URLs que apunten a direcciones IP internas",
                    )
        except socket.gaierror:
            raise HTTPException(status_code=400, detail=f"No se puede resolver el hostname: {hostname}")

    async def _validate_pdf_url(url: str) -> None:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"}:
            raise HTTPException(status_code=400, detail="file_url debe usar http/https")

        _check_ssrf(url)

        try:
            async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as client:
                head = await client.head(url)
                if head.status_code >= 400:
                    raise HTTPException(
                        status_code=400,
                        detail=f"No se puede acceder al archivo URL (HEAD {head.status_code}).",
                    )

                content_type = (head.headers.get("content-type") or "").lower()
                if content_type and "pdf" not in content_type:
                    raise HTTPException(
                        status_code=400,
                        detail=f"El content-type no parece PDF: {content_type}",
                    )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error validando URL: {e}")

    async def _download_pdf_with_retry(url: str) -> bytes:
        delays = [1, 2, 4]
        last_error: Exception | None = None

        for idx, delay in enumerate(delays, 1):
            try:
                async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                    response = await client.get(url)
                    response.raise_for_status()

                payload = response.content
                if not payload:
                    raise ValueError("El archivo descargado está vacío")
                if len(payload) > 100 * 1024 * 1024:
                    raise ValueError("File too large (max 100 MB)")
                return payload
            except Exception as e:
                last_error = e
                if idx == len(delays):
                    break
                await asyncio.sleep(delay)

        raise HTTPException(
            status_code=502,
            detail=f"No se pudo descargar el PDF tras {len(delays)} intentos: {last_error}",
        )

    file_bytes: bytes
    ingest_filename: str

    if storage_path:
        # ── Storage mode: download from Supabase Storage ──
        if not filename:
            # Derive filename from storage_path
            filename = os.path.basename(storage_path)
        ingest_filename = filename

        try:
            from supabase._async.client import create_client as acreate_client

            sb = await acreate_client(
                _config.supabase_url,
                _config.supabase_service_key,
            )
            file_bytes = await sb.storage.from_("documentos").download(storage_path)
            logger.info(
                "Descargado desde Storage: %s (%d bytes)", storage_path, len(file_bytes)
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error descargando de Supabase Storage ({storage_path}): {e}",
            )
    elif file_url:
        await _validate_pdf_url(file_url)
        file_bytes = await _download_pdf_with_retry(file_url)

        parsed = urlparse(file_url)
        inferred_name = os.path.basename(parsed.path) or "documento.pdf"
        ingest_filename = filename or inferred_name
        if not ingest_filename.lower().endswith(".pdf"):
            ingest_filename = f"{ingest_filename}.pdf"
    else:
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename is required")
        ingest_filename = file.filename
        file_bytes = await file.read()

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="File is empty")

    if len(file_bytes) > 100 * 1024 * 1024:  # 100 MB limit
        raise HTTPException(status_code=413, detail="File too large (max 100 MB)")

    try:
        result = await service.ingest(
            file_bytes=file_bytes,
            filename=ingest_filename,
            project_id=project_id,
            rag_scope=rag_scope,
            password=password,
        )
        if not result.success:
            raise HTTPException(
                status_code=422,
                detail=result.error or f"Error al procesar '{ingest_filename}'.",
            )
        return result.to_dict()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# RAG QUERY - Consulta con respuesta generada por LLM
# ═══════════════════════════════════════════════════════════════

class RAGQueryRequest(BaseModel):
    query: str = Field(min_length=1, max_length=10000)
    project_id: Optional[str] = None
    scope: Optional[str] = None  # "general", "project", or None (both)
    top_k: int = Field(default=5, ge=1, le=50)


class RAGQueryResponse(BaseModel):
    answer: str
    sources: list[dict]
    query: str
    scope_used: list[str]


@app.post("/api/rag/query", response_model=RAGQueryResponse)
@limiter.limit("30/minute")
async def rag_query(request: Request, payload: RAGQueryRequest):
    """
    Consulta al RAG de documentos normativos y técnicos.
    Busca chunks relevantes y genera una respuesta con Claude.
    """
    if rag_service is None or _config is None:
        raise HTTPException(status_code=503, detail="RAG service not initialized")

    # Determinar scopes a consultar
    scopes = None
    if payload.scope == "general":
        scopes = [RAGScope.GENERAL]
    elif payload.scope == "project":
        scopes = [RAGScope.PROJECT]

    try:
        # Buscar en el RAG
        rag_response = await rag_service.search(
            query=payload.query,
            project_id=payload.project_id,
            scopes=scopes,
            top_k_per_scope=payload.top_k,
        )

        # Generar respuesta con Claude usando el contexto recuperado
        from anthropic import AsyncAnthropic
        claude = AsyncAnthropic(api_key=_config.anthropic_api_key, max_retries=4)

        system_prompt = (
            "Eres un experto en gestión de residuos industriales en España. "
            "Respondes preguntas basándote ESTRICTAMENTE en el contexto proporcionado. "
            "Si el contexto no contiene información suficiente, lo indicas claramente. "
            "Cita las fuentes de tu respuesta (nombre del documento y tipo). "
            "Responde siempre en español."
        )

        user_prompt = (
            f"{rag_response.context_text}\n\n"
            f"Basándote en el contexto anterior, responde a esta pregunta:\n"
            f"{payload.query}"
        )

        try:
            message = await claude.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            answer = message.content[0].text
        except Exception as claude_err:
            is_overloaded = "overloaded" in str(claude_err).lower()
            if is_overloaded and _config.gemini_api_key:
                logger.warning("Claude overloaded in rag/query, falling back to Gemini 2.5 Pro")
                from google import genai
                from google.genai import types as genai_types
                gemini_client = genai.Client(api_key=_config.gemini_api_key)
                response = await gemini_client.aio.models.generate_content(
                    model="gemini-2.5-pro",
                    contents=[user_prompt],
                    config=genai_types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        max_output_tokens=2000,
                        temperature=0.7,
                    ),
                )
                answer = response.text or ""
            else:
                raise

        # Preparar fuentes
        sources = [
            {
                "document_id": r.document_id,
                "title": r.doc_title,
                "doc_type": r.doc_type,
                "chunk_type": r.chunk_type,
                "similarity": round(r.similarity, 3),
                "scope": r.rag_scope.value if isinstance(r.rag_scope, RAGScope) else r.rag_scope,
                "excerpt": r.content[:200] + "..." if len(r.content) > 200 else r.content,
            }
            for r in rag_response.results
        ]

        scope_used = []
        if rag_response.general_results:
            scope_used.append("general")
        if rag_response.project_results:
            scope_used.append("project")

        return RAGQueryResponse(
            answer=answer,
            sources=sources,
            query=payload.query,
            scope_used=scope_used,
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# ASESOR IA - Consultor experto en gestión de residuos
# ═══════════════════════════════════════════════════════════════

ADVISOR_SYSTEM_PROMPT = """Eres un asesor experto senior en gestión de residuos industriales en España, con más de 20 años de experiencia. Tu perfil profesional abarca el dominio completo de la legislación española y europea de residuos (Ley 7/2022, RD 553/2020, Directiva 2008/98/CE, Reglamento CLP, ADR), la clasificación LER en profundidad, las 15 propiedades de peligrosidad HP1-HP15 según el Reglamento (UE) 1357/2014, los BREFs sectoriales, estrategias de desclasificación y valorización, autorizaciones ambientales integradas (AAI), DARI, registro de producción, y precios de mercado de gestión por tipo y zona.

## FORMATO DE RESPUESTA — REGLA PRINCIPAL

Escribe como un consultor senior redactando un informe de consultoría. NUNCA uses formato epistolar (nada de "Estimado", "Atentamente", saludos ni despedidas). Tus respuestas deben ser PROSA NARRATIVA: párrafos desarrollados que expliquen el razonamiento, conecten los conceptos entre sí y desarrollen las implicaciones prácticas de cada punto. Este es tu modo por defecto, siempre.

Las listas de viñetas (bullets) y las listas numeradas NO son tu formato principal. Solo úsalas en estos casos muy concretos: enumerar los pasos secuenciales de un procedimiento administrativo, listar requisitos documentales específicos, o presentar un checklist de verificación. Fuera de esos tres casos, escribe siempre en párrafos narrativos.

Las tablas son valiosas cuando necesitas comparar alternativas (coste, plazo, pros/contras), presentar datos numéricos, o contrastar opciones de gestión. Úsalas cuando aporten claridad visual que un párrafo no puede dar.

Los encabezados (## y ###) sirven para organizar secciones temáticas en respuestas largas. Dentro de cada sección, desarrolla párrafos completos con razonamiento conectado.

Calibra la extensión a la complejidad de la pregunta. Una duda puntual merece una respuesta precisa y directa en uno o dos párrafos. Una consulta técnica compleja merece un análisis exhaustivo con contexto normativo, matices técnicos y recomendaciones detalladas.

EJEMPLO DE LO QUE NUNCA DEBES HACER (formato ficha/esquemático):
"- Código LER: 170503*
- Peligrosidad: HP14 Ecotóxico
- Gestión: vertedero de seguridad
- Coste: 80-120 €/t"

EJEMPLO DEL ESTILO CORRECTO (prosa de consultoría):
"Este residuo se clasifica bajo el código LER 170503* (tierras y piedras que contienen sustancias peligrosas), un código espejo cuya peligrosidad depende de la concentración de contaminantes presentes. El análisis muestra concentraciones de hidrocarburos totales de petróleo (TPH) de 3.200 mg/kg, lo que supera ampliamente el umbral de 1.000 mg/kg establecido para la propiedad HP14 (Ecotóxico) según el Reglamento 1357/2014. Esto implica que el residuo debe gestionarse obligatoriamente como peligroso, con destino a un depósito de seguridad autorizado. El coste orientativo de esta gestión oscila entre 80 y 120 €/tonelada dependiendo de la zona geográfica y el gestor, aunque conviene solicitar al menos tres ofertas dado que el mercado presenta variaciones significativas entre operadores."

## CÓMO RESPONDER SEGÚN EL TIPO DE CONSULTA

Cuando analices un residuo, desarrolla un dictamen técnico integrado: justifica la clasificación LER explicando por qué corresponde ese código y no otro, determina la peligrosidad analizando cada propiedad HP relevante con las sustancias y concentraciones que la determinan, y presenta las opciones de gestión (valorización, tratamiento, eliminación) con costes orientativos y una recomendación fundamentada.

Cuando te pregunten sobre desclasificación, redacta un informe que desarrolle el razonamiento completo: qué propiedades HP hay que eliminar y por qué, qué tratamientos existen con sus ventajas e inconvenientes prácticos, qué análisis de laboratorio se necesitan para demostrar la desclasificación, y cuál es el procedimiento administrativo ante la autoridad competente.

Cuando cites normativa, hazlo siempre en contexto narrativo. No basta con escribir "artículo 20 de la Ley 7/2022"; debes explicar qué establece ese artículo y cómo afecta al caso concreto del cliente en su operativa diaria.

Si tienes contexto del RAG, úsalo como fuente principal y complementa con tu conocimiento experto, extrayendo todos los datos relevantes y poniéndolos en relación con la consulta. Si no tienes contexto del RAG, responde con tu expertise y deja claro que no has encontrado documentos específicos en la base de conocimiento.

Si el usuario sube un análisis químico, redacta una interpretación profesional que ponga en contexto el análisis, interprete cada valor relevante frente a sus límites legales, determine códigos LER y propiedades HP con su razonamiento explícito, y recomiende acciones de gestión fundamentadas.

Sé siempre concreto y técnico (códigos LER exactos, artículos de ley, concentraciones límite, propiedades HP), pero integra esos datos dentro de tu análisis narrativo en lugar de presentarlos como fichas aisladas.

## CONOCIMIENTO TÉCNICO

Tu expertise abarca la clasificación de residuos (LER, códigos espejo, peligrosidad), las 15 propiedades HP desde HP1 Explosivo hasta HP15 Residuo con peligrosidad diferida, estrategias de desclasificación y valorización, obligaciones legales del productor y poseedor, contratos con gestores autorizados, DARI y registro cronológico, almacenamiento temporal con sus límites y condiciones, transporte ADR de mercancías peligrosas, MTD/BAT (Mejores Técnicas Disponibles) por sector, y economía circular incluyendo simbiosis industrial.

## RAZONAMIENTO INTERNO

Ante consultas complejas, antes de redactar tu respuesta: identifica el tipo de consulta, recopila todos los datos relevantes del contexto RAG, documentos adjuntos y tu expertise propio, aplica la legislación y criterios técnicos correspondientes citando artículos exactos, analiza alternativas cuando existan, y estructura tu respuesta con recomendaciones concretas y accionables. Cita siempre las fuentes normativas específicas (artículo, anexo, ley, real decreto).

## CITACIÓN DE FUENTES DOCUMENTALES

Cuando tu respuesta se base en documentos del contexto RAG (tanto del proyecto como de la base de conocimiento general), integra las referencias de forma natural en tu texto narrativo. No te limites a usar la información: indica de dónde procede para que el consultor pueda verificarla y trazarla. Por ejemplo: "De acuerdo con el BREF de tratamiento de residuos recogido en la base de conocimiento..." o "Según consta en la Autorización Ambiental Integrada del proyecto...". Cuando cites legislación extraída de un documento RAG, menciona tanto la norma como el documento fuente. Esto es especialmente importante cuando los datos provienen de documentos del proyecto (facturas, contratos, AAI, registro de producción), ya que el consultor necesita saber exactamente de qué documento se extrae cada dato.

## BÚSQUEDA WEB

Tienes acceso a búsqueda web como complemento del RAG. Úsala cuando aporte valor real:
- Verificar vigencia de normativa citada en el RAG (BOE, DOUE, transposiciones recientes).
- Consultas sobre gestores autorizados, plantas de tratamiento o instalaciones específicas.
- Confirmar umbrales, concentraciones límite, valores técnicos o clasificaciones HP actuales.
- Precios de mercado, tasas o cánones de gestión de residuos.
No la uses para preguntas que puedes responder bien con el RAG y tu conocimiento experto. Cuando uses resultados web, indica la fuente.

IMPORTANTE: Tu objetivo principal es RESPONDER al consultor con un análisis completo y profesional. Las búsquedas web complementan tu respuesta, nunca la sustituyen. Siempre genera tu respuesta de texto, aunque no encuentres resultados web relevantes.

Responde siempre en español."""


ADVISOR_REPORT_MODE_ADDENDUM = """

## MODO INFORME PROFESIONAL (ACTIVADO)

Cuando la consulta lo requiera, responde como entregable de consultoria medioambiental en gestion de residuos.
Usa exactamente estos encabezados de nivel 2 (##), en este orden:
1) ## 1. Resumen ejecutivo
2) ## 2. Alcance, metodologia y limitaciones
3) ## 3. Contexto operativo y linea base
4) ## 4. Evaluacion de cumplimiento normativo
5) ## 5. Analisis economico de la gestion de residuos
6) ## 6. Oportunidades de mejora y eficiencia
7) ## 7. Plan de accion priorizado (30-60-90 dias)
8) ## 8. Matriz de riesgos y recomendaciones de control
9) ## 9. Conclusion ejecutiva
10) ## 10. Anexo de trazabilidad tecnica

Reglas de este modo:
- No inventes datos; marca cualquier ausencia como "limitacion de evidencia".
- Integra datos cuantitativos disponibles (EUR, t/ano, LER, plazos, gestores, norma aplicable).
- En el plan 30-60-90 incluye responsable sugerido, dependencia y resultado esperado por accion.
- En matriz de riesgos agrupa al menos: legal, operativo y economico.
"""


def _is_professional_report_mode(query: str, analysis_context: Optional[dict]) -> bool:
    """Detects whether the advisor should answer in consultancy report format."""
    if analysis_context:
        mode = str(analysis_context.get("response_mode", "")).lower().strip()
        output_format = str(analysis_context.get("output_format", "")).lower().strip()
        if mode in {"report", "professional_report", "consulting_report"}:
            return True
        if output_format in {"report", "professional_report", "consulting_report"}:
            return True
        if analysis_context.get("force_professional_report") is True:
            return True

    q = (query or "").lower()
    trigger_phrases = [
        "informe",
        "informe profesional",
        "informe ejecutivo",
        "formato consultoria",
        "estandar de consultoria",
        "dictamen tecnico",
    ]
    return any(t in q for t in trigger_phrases)


def _build_analysis_context_addendum(ctx: dict) -> str:
    """Build a system prompt addendum with the HITL analysis context."""
    parts = ["\n\n## CONTEXTO DEL ANALISIS EN CURSO"]
    phase = ctx.get("phase", "")
    project_name = ctx.get("projectName", "")

    if project_name:
        parts.append(f"Estas asistiendo a un consultor que analiza el proyecto **{project_name}**.")

    if phase == "plan_review":
        parts.append("El consultor esta revisando el PLAN DE ANALISIS propuesto por el coordinador IA.")
        parts.append("Ayudale a decidir que agentes activar, que foco dar a cada uno, y que instrucciones escribir.")

        plan = ctx.get("plan", {})
        if plan:
            agents = plan.get("agents", [])
            if agents:
                parts.append("\n### Agentes propuestos:")
                for a in agents:
                    status = "ACTIVADO" if a.get("enabled") else "desactivado"
                    parts.append(f"- **{a.get('id', '?')}** [{status}]: {a.get('reason', '')}")
                    if a.get("focus"):
                        parts.append(f"  Foco sugerido: {a['focus']}")

            gaps = plan.get("data_gaps", [])
            if gaps:
                parts.append("\n### Carencias de datos:")
                for g in gaps:
                    parts.append(f"- {g}")

            summary = plan.get("data_summary", {})
            if summary:
                parts.append(f"\n### Datos del proyecto: {summary.get('total_documents', 0)} docs, "
                             f"{summary.get('inventory_items', 0)} residuos, "
                             f"{summary.get('contracts', 0)} contratos, "
                             f"{summary.get('invoice_lines', 0)} lineas factura")

    elif phase == "results_review":
        parts.append("El consultor esta revisando los RESULTADOS del analisis y decidiendo si lanzar una 2a vuelta.")
        parts.append("Ayudale a interpretar los hallazgos y decidir que profundizar.")

        findings = ctx.get("findings", [])
        if findings:
            critical = [f for f in findings if f.get("severidad") == "critica"]
            high = [f for f in findings if f.get("severidad") == "alta"]
            parts.append(f"\n### Hallazgos: {len(findings)} total, {len(critical)} criticos, {len(high)} altos")

            for f in (critical + high)[:10]:
                ahorro = f.get("ahorro_eur_ano", 0)
                ahorro_str = f" ({ahorro:,.0f} EUR/a)" if ahorro else ""
                parts.append(f"- [{f.get('severidad', '?').upper()}] [{f.get('agente', '?')}] "
                             f"{f.get('descripcion', '')}{ahorro_str}")

    parts.append("\nResponde en el contexto de este analisis. Se concreto y util para las decisiones del consultor.")
    return "\n".join(parts)


class AdvisorMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class FileAttachment(BaseModel):
    name: str = Field(max_length=500)
    type: str  # "image", "document", or "binary"
    content: str  # base64 for images/binaries, extracted text for documents
    mime_type: Optional[str] = None  # e.g., "image/png", "application/pdf"
    size: int = 0


class AdvisorRequest(BaseModel):
    query: str = Field(min_length=1, max_length=20000)
    conversation_history: list[AdvisorMessage] = []
    project_id: Optional[str] = None
    # Multi-file support (up to 6)
    files: Optional[list[FileAttachment]] = Field(default=None, max_length=6)
    urls: Optional[list[str]] = Field(default=None, max_length=6)
    # Optional agentic folder scan in Google Drive
    consultant_id: Optional[str] = None
    gdrive_folder_id: Optional[str] = None
    gdrive_max_files: int = Field(default=12, ge=1, le=30)
    # HITL: analysis context when advisor is embedded in plan review or results
    analysis_context: Optional[dict] = None
    # Google Drive folder context (ephemeral, not persisted)
    drive_context: Optional[str] = None
    drive_files: Optional[list[dict]] = None
    # Legacy single-file support (backward compatibility)
    file_content: Optional[str] = None
    file_name: Optional[str] = None
    # Model selection
    model_override: Optional[str] = None   # Override: 'claude-opus-4-6', 'gpt-5.2', etc.
    tier: Optional[str] = None              # 'standard' or 'pro_plus'

    @field_validator("tier")
    @classmethod
    def validate_tier(cls, v):
        if v is not None and v not in _VALID_TIERS:
            raise ValueError(f"tier debe ser 'standard' o 'pro_plus', recibido: {v!r}")
        return v

    @field_validator("model_override")
    @classmethod
    def validate_model(cls, v):
        if v is not None and v not in _VALID_MODELS:
            raise ValueError(f"model_override no reconocido: {v!r}")
        return v


class AdvisorResponse(BaseModel):
    answer: str
    sources: list[dict]
    rag_context_used: bool


# ─── Server-side file text extraction ─────────────────────────────

def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text + tables from a PDF. Tries pdfplumber first, falls back to pdfminer."""
    import io

    # Try pdfplumber (best quality: text + tables)
    try:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages[:50], 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(f"--- Pagina {i} ---\n{page_text}")

                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    rows = []
                    for row in table:
                        cells = [str(cell or "").strip() for cell in row]
                        rows.append(" | ".join(cells))
                    if rows:
                        parts.append(f"[Tabla pagina {i}]\n" + "\n".join(rows))

        if parts:
            return "\n\n".join(parts)
    except ImportError:
        pass
    except Exception as e:
        logger.warning("pdfplumber failed for PDF: %s, trying fallback", e)

    # Fallback: pdfminer (text only, no tables)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(file_bytes), maxpages=50)
        if text and text.strip():
            return text.strip()
    except ImportError:
        pass
    except Exception as e:
        logger.warning("pdfminer fallback also failed: %s", e)

    return "[PDF: no se pudo extraer texto. El archivo puede estar escaneado o protegido.]"


def _extract_excel_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from Excel/CSV using pandas + openpyxl."""
    import pandas as pd
    import io

    parts: list[str] = []
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "csv":
            # Auto-detect separator
            sample = file_bytes[:4096].decode("utf-8", errors="replace")
            sep = ";" if sample.count(";") > sample.count(",") else ","
            df = pd.read_csv(io.BytesIO(file_bytes), sep=sep, on_bad_lines="skip")
            parts.append(f"--- CSV ({len(df)} filas x {len(df.columns)} columnas) ---")
            parts.append(df.to_markdown(index=False))
        else:
            xl = pd.ExcelFile(io.BytesIO(file_bytes))
            for sheet_name in xl.sheet_names[:10]:  # Max 10 sheets
                df = pd.read_excel(xl, sheet_name=sheet_name)
                if df.empty:
                    continue
                parts.append(
                    f"--- Hoja: {sheet_name} ({len(df)} filas x {len(df.columns)} columnas) ---"
                )
                # Limit rows for very large sheets
                if len(df) > 200:
                    parts.append(df.head(200).to_markdown(index=False))
                    parts.append(f"... ({len(df) - 200} filas mas omitidas)")
                else:
                    parts.append(df.to_markdown(index=False))
    except Exception as e:
        parts.append(f"[Error extrayendo Excel/CSV: {e}]")

    return "\n\n".join(parts) if parts else "[Archivo Excel/CSV sin datos]"


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from Word DOCX using python-docx."""
    import io

    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n[Tabla {t_idx + 1}]\n" + "\n".join(rows))

        return "\n".join(parts) if parts else "[Documento Word sin contenido]"
    except Exception as e:
        return f"[Error extrayendo Word: {e}]"


def _extract_binary_text(file_bytes: bytes, filename: str) -> str:
    """Route binary file to the appropriate text extractor."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf_text(file_bytes)
    elif ext in ("xlsx", "xls", "xlsm"):
        return _extract_excel_text(file_bytes, filename)
    elif ext == "csv":
        return _extract_excel_text(file_bytes, filename)
    elif ext in ("docx", "doc"):
        return _extract_docx_text(file_bytes)
    else:
        # Fallback: try reading as text
        try:
            return file_bytes.decode("utf-8", errors="replace")[:15000]
        except Exception:
            return f"[Formato no soportado: .{ext}]"


# ─── URL content fetching ─────────────────────────────────────────

async def _fetch_url_content(url: str) -> str:
    """Fetch text content from a URL for advisor context."""
    import httpx

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as client:
            resp = await client.get(url, headers={
                "User-Agent": "ResidusIA-Advisor/1.0",
                "Accept": "text/html,text/plain,application/json,*/*",
            })
            resp.raise_for_status()

            content_type = resp.headers.get("content-type", "")
            text = resp.text

            # Strip HTML tags for a rough text extraction
            if "html" in content_type:
                import re
                # Remove script/style blocks
                text = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', text, flags=re.DOTALL | re.IGNORECASE)
                # Remove HTML tags
                text = re.sub(r'<[^>]+>', ' ', text)
                # Collapse whitespace
                text = re.sub(r'\s+', ' ', text).strip()

            return text[:15000]
    except Exception as e:
        return f"[Error al obtener URL {url}: {e}]"


# ─── Advisor core logic (shared by JSON and FormData endpoints) ───

IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff", "tif"}


async def _collect_gdrive_docs_for_advisor(
    consultant_id: str,
    folder_id: str,
    max_files: int = 12,
) -> tuple[list[tuple[str, str]], list[str]]:
    """Read files from a Drive folder (read-only) for advisor analysis context."""
    docs: list[tuple[str, str]] = []
    warnings: list[str] = []

    gd, _ = await _get_gdrive_service(consultant_id)

    supported_extensions = {
        ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt", ".html", ".htm", ".md"
    }

    all_files = await asyncio.to_thread(
        gd.list_all_files_recursive,
        folder_id,
        supported_extensions,
    )

    if not all_files:
        return docs, ["[Google Drive] La carpeta seleccionada no contiene documentos compatibles."]

    file_limit = max(1, min(max_files, 30))

    for f in all_files[:file_limit]:
        file_id = f.get("id", "")
        file_name = f.get("name", "archivo")
        file_path = f.get("path", "")
        if not file_id:
            continue

        try:
            file_bytes, downloaded_name, _mime = await asyncio.to_thread(gd.download_file, file_id)
            effective_name = downloaded_name or file_name
            extracted = await asyncio.to_thread(_extract_binary_text, file_bytes, effective_name)
            prefixed_name = f"[GD] {file_path}" if file_path else f"[GD] {effective_name}"
            docs.append((prefixed_name, extracted[:15000]))
        except Exception as e:
            warnings.append(f"[Google Drive] No se pudo leer '{file_name}': {e}")

    if len(all_files) > file_limit:
        warnings.append(
            f"[Google Drive] Se analizaron {file_limit} archivos de {len(all_files)} disponibles (limite configurado)."
        )

    return docs, warnings


async def _run_advisor(
    query: str,
    conversation_history: list[dict],
    project_id: Optional[str],
    processed_docs: list[tuple[str, str]],
    image_blocks: list[dict],
    url_list: list[str],
    analysis_context: Optional[dict] = None,
    consultant_id: Optional[str] = None,
    model_override: Optional[str] = None,
    tier: Optional[str] = None,
) -> dict:
    """
    Core advisor logic: RAG search → build prompt → Claude with thinking.
    Returns {"answer": str, "sources": list, "rag_context_used": bool}.
    """
    if rag_service is None or _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    from anthropic import AsyncAnthropic

    # 1. RAG search
    scopes = [RAGScope.GENERAL]
    if project_id:
        scopes.append(RAGScope.PROJECT)

    rag_response = await rag_service.search(
        query=query,
        project_id=project_id,
        scopes=scopes,
        top_k_per_scope=12,
        similarity_threshold=0.50,
    )
    has_rag_context = bool(rag_response.results)

    # 2. Fetch URLs in parallel
    url_contents: list[tuple[str, str]] = []
    if url_list:
        tasks = [_fetch_url_content(u) for u in url_list[:6]]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        for url, result in zip(url_list, results):
            if isinstance(result, Exception):
                url_contents.append((url, f"[Error al obtener URL: {result}]"))
            else:
                url_contents.append((url, result))

    # 3. Build text context
    text_parts: list[str] = []

    if has_rag_context:
        text_parts.append(
            "CONTEXTO DE LA BASE DE CONOCIMIENTO:\n"
            f"{rag_response.context_text}\n"
        )

    for i, (name, text) in enumerate(processed_docs, 1):
        text_parts.append(f"DOCUMENTO ADJUNTO {i} ({name}):\n{text}\n")

    for i, (url, content) in enumerate(url_contents, 1):
        text_parts.append(f"CONTENIDO DE URL {i} ({url}):\n{content}\n")

    if image_blocks:
        text_parts.append(
            f"Se han adjuntado {len(image_blocks)} imagen(es). "
            "Analiza cada imagen en detalle: identifica residuos, codigos, "
            "etiquetas, valores de analisis quimicos, fichas de seguridad, "
            "o cualquier informacion relevante para la gestion de residuos.\n"
        )

    text_parts.append(f"PREGUNTA DEL CONSULTOR:\n{query}")

    # 4. Build multimodal content blocks
    user_content_blocks: list[dict] = []
    for img in image_blocks:
        user_content_blocks.append(img)
    user_content_blocks.append({
        "type": "text",
        "text": "\n---\n".join(text_parts),
    })

    # 5. Build messages (truncate old assistant responses to save tokens)
    messages = _truncate_history(conversation_history)
    messages.append({"role": "user", "content": user_content_blocks})

    # 6. Build system prompt, injecting analysis context if available
    system_prompt = ADVISOR_SYSTEM_PROMPT
    if _is_professional_report_mode(query, analysis_context):
        system_prompt += ADVISOR_REPORT_MODE_ADDENDUM
    if analysis_context:
        system_prompt += _build_analysis_context_addendum(analysis_context)

    # Adaptive thinking: less budget on follow-ups (context already established)
    thinking_budget = 10000 if conversation_history else 24000

    # 7. Resolve model via ModelRouter chain
    effective_tier = tier or "standard"
    chain = await _model_router.get_consultant_chain(
        service="advisor",
        consultant_id=consultant_id,
        tier_override=effective_tier,
        model_override=model_override,
    )

    answer = ""
    web_sources: list[dict] = []
    model_used = "claude-sonnet-4"
    input_tokens = 0
    output_tokens = 0
    call_start = time.monotonic()

    # Web search tool (Anthropic-only, Gemini/OpenAI have their own)
    web_search_tool = {
        "type": "web_search_20250305",
        "name": "web_search",
        "max_uses": 2,
    }

    # Try each model in the fallback chain
    last_error = None
    for i, model_id in enumerate(chain):
        provider = MODEL_PROVIDERS.get(model_id, "unknown")

        # Cost Guard check
        if _cost_guard and consultant_id:
            check = await _cost_guard.check(provider, consultant_id)
            if not check.allowed:
                logger.info("CostGuard blocked %s for advisor: %s", model_id, check.reason)
                continue

        try:
            if provider == "anthropic":
                from anthropic import AsyncAnthropic
                claude = AsyncAnthropic(api_key=_config.anthropic_api_key, max_retries=4)
                api_model = MODEL_API_IDS.get(model_id, model_id)

                async with claude.messages.stream(
                    model=api_model,
                    max_tokens=32000,
                    thinking={
                        "type": "enabled",
                        "budget_tokens": thinking_budget,
                    },
                    tools=[web_search_tool],
                    system=system_prompt,
                    messages=messages,
                ) as stream:
                    final_response = await stream.get_final_message()

                input_tokens = final_response.usage.input_tokens
                output_tokens = final_response.usage.output_tokens

                for block in final_response.content:
                    if block.type == "text":
                        answer = block.text
                    elif block.type == "web_search_tool_result":
                        for item in getattr(block, "content", []):
                            if getattr(item, "type", None) == "web_search_result":
                                web_sources.append({
                                    "title": getattr(item, "title", ""),
                                    "url": getattr(item, "url", ""),
                                    "scope": "web",
                                })

            elif provider == "openai":
                # Responses API con web_search_preview
                response = await _model_router.call_openai(
                    model_id, system_prompt, messages,
                    max_tokens=32000, web_search=True,
                )
                from pipeline.model_router import (
                    _extract_openai_responses_text,
                    _extract_openai_web_sources,
                )
                answer = _extract_openai_responses_text(response)
                oai_ws = _extract_openai_web_sources(response)
                for ws in oai_ws:
                    ws["scope"] = "web"
                web_sources.extend(oai_ws)
                usage = getattr(response, "usage", None)
                input_tokens = getattr(usage, "input_tokens", 0) if usage else 0
                output_tokens = getattr(usage, "output_tokens", 0) if usage else 0

            elif provider == "google":
                # Google con thinking budget + Google Search grounding
                response = await _model_router.call_google(
                    model_id, system_prompt, messages,
                    max_tokens=32000,
                    thinking_budget=thinking_budget,
                    web_search=True,
                )
                answer = response.text or ""
                from pipeline.model_router import _extract_google_web_sources
                google_ws = _extract_google_web_sources(response)
                for ws in google_ws:
                    ws["scope"] = "web"
                web_sources.extend(google_ws)
                usage_meta = getattr(response, "usage_metadata", None)
                input_tokens = getattr(usage_meta, "prompt_token_count", 0) if usage_meta else 0
                output_tokens = getattr(usage_meta, "candidates_token_count", 0) if usage_meta else 0
            else:
                continue

            model_used = model_id
            last_error = None
            break  # success

        except Exception as e:
            last_error = e
            is_overloaded = "overloaded" in str(e).lower()
            logger.warning("Advisor: %s failed (overloaded=%s): %s", model_id, is_overloaded, str(e)[:200])
            continue

    if last_error and not answer:
        raise last_error

    call_duration = int((time.monotonic() - call_start) * 1000)

    # 8. Record cost
    if _cost_guard:
        await _cost_guard.record(
            model=model_used, service="advisor", operation="advisor_chat",
            input_tokens=input_tokens, output_tokens=output_tokens,
            duration_ms=call_duration, consultant_id=consultant_id,
            project_id=project_id,
            metadata={"web_searches": len(web_sources), "tier": effective_tier,
                       "thinking_budget": thinking_budget, "chain": chain},
        )

    # 8. Combine RAG sources + web sources (deduplicated by document)
    sources = _deduplicate_sources(rag_response.results)

    # Add web sources (deduplicated by URL)
    seen_urls: set[str] = set()
    for ws in web_sources:
        if ws["url"] and ws["url"] not in seen_urls:
            seen_urls.add(ws["url"])
            sources.append({
                "document_id": ws["url"],
                "title": ws["title"],
                "doc_type": "web",
                "similarity": 0,
                "scope": "web",
                "excerpt": ws["url"],
            })

    web_search_used = len(web_sources) > 0
    cost = calculate_cost(model_used, input_tokens, output_tokens)
    logger.info(
        "Advisor: model=%s, RAG=%s, web_search=%s (%d results), docs=%d, images=%d, "
        "tokens=%d in + %d out, cost=$%.4f, duration=%dms",
        model_used, has_rag_context, web_search_used, len(web_sources),
        len(processed_docs), len(image_blocks),
        input_tokens, output_tokens, cost, call_duration,
    )

    return {
        "answer": answer,
        "sources": sources,
        "rag_context_used": has_rag_context,
        "web_search_used": web_search_used,
        "model_used": model_used,
        "cost_usd": cost,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
    }


# ─── Advisor helpers ─────────────────────────────────────────────


def _deduplicate_sources(results) -> list[dict]:
    """Deduplicate RAG sources by document_id, keeping the highest similarity."""
    seen: dict[str, dict] = {}
    for r in results:
        doc_id = r.document_id
        similarity = round(r.similarity, 3)
        scope = r.rag_scope.value if isinstance(r.rag_scope, RAGScope) else r.rag_scope
        if doc_id not in seen or similarity > seen[doc_id]["similarity"]:
            seen[doc_id] = {
                "document_id": doc_id,
                "title": r.doc_title,
                "doc_type": r.doc_type,
                "similarity": similarity,
                "scope": scope,
                "excerpt": r.content[:200] + "..." if len(r.content) > 200 else r.content,
            }
    return list(seen.values())


MAX_HISTORY_MSG_CHARS = 1500


def _truncate_history(history: list[dict], max_msgs: int = 10) -> list[dict]:
    """Truncate old assistant messages to reduce token count on follow-ups."""
    truncated = []
    for msg in history[-max_msgs:]:
        content = msg.get("content", "")
        if msg.get("role") == "assistant" and len(content) > MAX_HISTORY_MSG_CHARS:
            content = content[:MAX_HISTORY_MSG_CHARS] + "\n\n[... respuesta anterior truncada por longitud ...]"
        truncated.append({"role": msg["role"], "content": content})
    return truncated


def _sse_event(event_type: str, data) -> str:
    """Format a Server-Sent Event."""
    return f"event: {event_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ─── Advisor endpoint: JSON (text-only, through Vercel proxy) ────

@app.post("/api/advisor")
@limiter.limit("20/minute")
async def advisor_query(request: Request, payload: AdvisorRequest):
    """
    Asesor IA - JSON endpoint (for text-only queries through Vercel proxy).
    For file uploads, use POST /api/advisor/chat with FormData.
    """
    import base64

    try:
        # Normalize files
        files = list(payload.files or [])
        if not files and payload.file_content:
            files.append(FileAttachment(
                name=payload.file_name or "archivo",
                type="document",
                content=payload.file_content[:15000],
            ))

        processed_docs: list[tuple[str, str]] = []
        image_blocks: list[dict] = []

        for f in files[:6]:
            if f.type == "image" and f.mime_type:
                image_blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": f.mime_type,
                        "data": f.content,
                    },
                })
            elif f.type == "binary":
                try:
                    file_bytes = base64.b64decode(f.content)
                    extracted = await asyncio.to_thread(
                        _extract_binary_text, file_bytes, f.name
                    )
                    processed_docs.append((f.name, extracted[:15000]))
                except Exception as e:
                    processed_docs.append((f.name, f"[Error procesando {f.name}: {e}]"))
            else:
                processed_docs.append((f.name, f.content[:15000]))

        if payload.consultant_id and payload.gdrive_folder_id:
            gd_docs, gd_warnings = await _collect_gdrive_docs_for_advisor(
                consultant_id=payload.consultant_id,
                folder_id=payload.gdrive_folder_id,
                max_files=payload.gdrive_max_files,
            )
            processed_docs.extend(gd_docs)
            for w in gd_warnings:
                processed_docs.append(("google_drive_notice", w))

        history = [{"role": m.role, "content": m.content} for m in payload.conversation_history]

        result = await _run_advisor(
            query=payload.query,
            conversation_history=history,
            project_id=payload.project_id,
            processed_docs=processed_docs,
            image_blocks=image_blocks,
            url_list=payload.urls or [],
            analysis_context=payload.analysis_context,
            consultant_id=payload.consultant_id,
            model_override=payload.model_override,
            tier=payload.tier,
        )
        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en advisor (JSON): {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─── Advisor endpoint: FormData (file uploads, direct from browser) ──

@app.post("/api/advisor/chat")
@limiter.limit("20/minute")
async def advisor_chat(
    request: Request,
    query: str = Form(...),
    conversation_history: str = Form(default="[]"),
    project_id: Optional[str] = Form(default=None),
    urls: str = Form(default="[]"),
    storage_files: str = Form(default="[]"),
    analysis_context: str = Form(default=""),
    consultant_id: Optional[str] = Form(default=None),
    gdrive_folder_id: Optional[str] = Form(default=None),
    gdrive_max_files: int = Form(default=12),
    model_override: Optional[str] = Form(default=None),
    tier: Optional[str] = Form(default=None),
    files: list[UploadFile] = File(default=[]),
):
    """
    Asesor IA - FormData endpoint for file uploads.
    Frontend calls this directly (bypassing Vercel's 4.5MB payload limit).
    Accepts real file uploads via multipart/form-data.
    Large files (>4MB) arrive as storage_files (JSON array of {name, type, storage_path})
    already uploaded to Supabase Storage.
    """
    import base64 as b64
    import json

    # Validar tier y model_override antes de procesar
    if tier and tier not in _VALID_TIERS:
        raise HTTPException(status_code=422, detail=f"tier invalido: {tier}")
    if model_override and model_override not in _VALID_MODELS:
        raise HTTPException(status_code=422, detail=f"model_override invalido: {model_override}")

    # Clamp gdrive_max_files
    gdrive_max_files = max(1, min(gdrive_max_files, 30))

    try:
        # Parse JSON fields
        try:
            history = json.loads(conversation_history)
        except (json.JSONDecodeError, TypeError):
            history = []

        try:
            url_list = json.loads(urls)
        except (json.JSONDecodeError, TypeError):
            url_list = []

        try:
            storage_file_list = json.loads(storage_files)
        except (json.JSONDecodeError, TypeError):
            storage_file_list = []

        # Process uploaded files
        processed_docs: list[tuple[str, str]] = []
        image_blocks: list[dict] = []

        # Process large files from Supabase Storage
        if storage_file_list and _config:
            from supabase._async.client import create_client as acreate_client

            sb = await acreate_client(
                _config.supabase_url,
                _config.supabase_service_key,
            )
            for sf in storage_file_list[:6]:
                sf_name = sf.get("name", "archivo")
                sf_path = sf.get("storage_path", "")
                if not sf_path:
                    continue
                try:
                    file_bytes = await sb.storage.from_("documentos").download(sf_path)
                    if len(file_bytes) == 0:
                        continue
                    ext = sf_name.rsplit(".", 1)[-1].lower() if "." in sf_name else ""
                    if ext in IMAGE_EXTENSIONS:
                        encoded = b64.b64encode(file_bytes).decode("ascii")
                        mime = sf.get("type") or f"image/{ext}"
                        image_blocks.append({
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": mime,
                                "data": encoded,
                            },
                        })
                        logger.info("Advisor: storage image %s (%d KB)", sf_name, len(file_bytes) // 1024)
                    else:
                        extracted = await asyncio.to_thread(
                            _extract_binary_text, file_bytes, sf_name
                        )
                        processed_docs.append((sf_name, extracted[:15000]))
                        logger.info("Advisor: storage doc %s (%d chars)", sf_name, len(extracted))
                except Exception as e:
                    processed_docs.append((sf_name, f"[Error descargando {sf_name} de Storage: {e}]"))
                    logger.warning("Advisor: storage download failed for %s: %s", sf_name, e)

        for upload in files[:6]:
            if not upload.filename:
                continue

            file_bytes = await upload.read()
            if len(file_bytes) == 0:
                continue
            if len(file_bytes) > 20 * 1024 * 1024:  # 20MB per file
                processed_docs.append(
                    (upload.filename, f"[Archivo demasiado grande: {upload.filename}]")
                )
                continue

            ext = upload.filename.rsplit(".", 1)[-1].lower() if "." in upload.filename else ""

            if ext in IMAGE_EXTENSIONS:
                # Images → base64 for Claude Vision
                encoded = b64.b64encode(file_bytes).decode("ascii")
                mime = upload.content_type or f"image/{ext}"
                image_blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": encoded,
                    },
                })
                logger.info("Advisor: image %s (%s, %d KB)", upload.filename, mime, len(file_bytes) // 1024)
            else:
                # Documents → extract text server-side
                try:
                    extracted = await asyncio.to_thread(
                        _extract_binary_text, file_bytes, upload.filename
                    )
                    processed_docs.append((upload.filename, extracted[:15000]))
                    logger.info(
                        "Advisor: extracted %d chars from %s",
                        len(extracted), upload.filename,
                    )
                except Exception as e:
                    processed_docs.append(
                        (upload.filename, f"[Error procesando {upload.filename}: {e}]")
                    )
                    logger.warning("Advisor: extraction failed for %s: %s", upload.filename, e)

        if consultant_id and gdrive_folder_id:
            gd_docs, gd_warnings = await _collect_gdrive_docs_for_advisor(
                consultant_id=consultant_id,
                folder_id=gdrive_folder_id,
                max_files=gdrive_max_files,
            )
            processed_docs.extend(gd_docs)
            for w in gd_warnings:
                processed_docs.append(("google_drive_notice", w))

        # Parse analysis_context
        parsed_analysis_context = None
        if analysis_context:
            try:
                parsed_analysis_context = json.loads(analysis_context)
            except (json.JSONDecodeError, TypeError):
                pass

        result = await _run_advisor(
            query=query,
            conversation_history=history,
            project_id=project_id,
            processed_docs=processed_docs,
            image_blocks=image_blocks,
            url_list=url_list,
            analysis_context=parsed_analysis_context,
            consultant_id=consultant_id,
            model_override=model_override,
            tier=tier,
        )
        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en advisor (FormData): {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─── Advisor endpoint: SSE streaming (no timeout) ────────────

@app.post("/api/advisor/stream")
@limiter.limit("20/minute")
async def advisor_stream(request: Request, payload: AdvisorRequest):
    """
    SSE streaming endpoint for the advisor.
    Returns Server-Sent Events instead of a single JSON response.
    This eliminates timeout issues because the connection stays alive
    as text is generated incrementally.

    Events emitted:
      - sources: RAG sources found (sent first)
      - text_delta: incremental text chunks
      - done: final metadata (web_search_used, web_sources)
      - error: if something fails
    """
    if rag_service is None or _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    async def generate():
        try:
            from anthropic import AsyncAnthropic

            # 1. RAG search
            scopes = [RAGScope.GENERAL]
            if payload.project_id:
                scopes.append(RAGScope.PROJECT)

            rag_response = await rag_service.search(
                query=payload.query,
                project_id=payload.project_id,
                scopes=scopes,
                top_k_per_scope=12,
                similarity_threshold=0.50,
            )
            has_rag_context = bool(rag_response.results)

            # Emit sources immediately (deduplicated by document)
            sources = _deduplicate_sources(rag_response.results)
            yield _sse_event("sources", {"sources": sources, "rag_context_used": has_rag_context})

            # 2. Build text context
            text_parts: list[str] = []
            if has_rag_context:
                text_parts.append(
                    "CONTEXTO DE LA BASE DE CONOCIMIENTO:\n"
                    f"{rag_response.context_text}\n"
                )
            if payload.drive_context:
                n_files = len(payload.drive_files) if payload.drive_files else 0
                text_parts.append(
                    f"DOCUMENTOS DE GOOGLE DRIVE ({n_files} archivos cargados como contexto):\n"
                    f"{payload.drive_context}\n"
                )
            text_parts.append(f"PREGUNTA DEL CONSULTOR:\n{payload.query}")

            user_content = "\n---\n".join(text_parts)

            # 3. Build messages with truncated history
            history = [{"role": m.role, "content": m.content} for m in payload.conversation_history]
            messages = _truncate_history(history)
            messages.append({"role": "user", "content": user_content})

            # 4. Adaptive thinking budget
            thinking_budget = 10000 if payload.conversation_history else 24000

            system_prompt = ADVISOR_SYSTEM_PROMPT
            if _is_professional_report_mode(payload.query, payload.analysis_context):
                system_prompt += ADVISOR_REPORT_MODE_ADDENDUM
            if payload.analysis_context:
                system_prompt += _build_analysis_context_addendum(payload.analysis_context)

            # 5. Resolve model chain via ModelRouter
            effective_tier = payload.tier or "standard"
            chain = await _model_router.get_consultant_chain(
                service="advisor",
                consultant_id=payload.consultant_id,
                tier_override=effective_tier,
                model_override=payload.model_override,
            )

            web_search_tool = {
                "type": "web_search_20250305",
                "name": "web_search",
                "max_uses": 2,
            }

            # 5b. Try each model in the fallback chain
            text_streamed = False
            model_used = chain[0] if chain else "claude-sonnet-4"
            final = None
            stream_input_tokens = 0
            stream_output_tokens = 0
            stream_start = time.monotonic()
            last_keepalive = stream_start
            gemini_web_sources: list[dict] = []
            oai_web_sources: list[dict] = []

            for chain_idx, model_id in enumerate(chain):
                provider = MODEL_PROVIDERS.get(model_id, "unknown")

                # Cost Guard check
                if _cost_guard and payload.consultant_id:
                    check = await _cost_guard.check(provider, payload.consultant_id)
                    if not check.allowed:
                        logger.info("CostGuard blocked %s for stream: %s", model_id, check.reason)
                        continue

                try:
                    if provider == "anthropic":
                        from anthropic import AsyncAnthropic
                        claude = AsyncAnthropic(api_key=_config.anthropic_api_key, max_retries=4)
                        api_model = MODEL_API_IDS.get(model_id, model_id)

                        async with claude.messages.stream(
                            model=api_model,
                            max_tokens=32000,
                            thinking={
                                "type": "enabled",
                                "budget_tokens": thinking_budget,
                            },
                            tools=[web_search_tool],
                            system=system_prompt,
                            messages=messages,
                        ) as stream:
                            async for event in stream:
                                if event.type == "text":
                                    text_streamed = True
                                    last_keepalive = time.monotonic()
                                    yield _sse_event("text_delta", {"text": event.text})
                                elif event.type == "content_block_start":
                                    block = getattr(event, "content_block", None)
                                    if block:
                                        btype = getattr(block, "type", None)
                                        if btype == "thinking":
                                            yield _sse_event("status", {"phase": "thinking"})
                                        elif btype in ("server_tool_use", "web_search_tool_result"):
                                            yield _sse_event("status", {"phase": "web_search"})
                                    last_keepalive = time.monotonic()
                                else:
                                    now = time.monotonic()
                                    if now - last_keepalive >= 5:
                                        yield ": keepalive\n\n"
                                        last_keepalive = now

                            final = await stream.get_final_message()

                        model_used = model_id
                        break  # success

                    elif provider == "google":
                        from google import genai
                        from google.genai import types as genai_types

                        if chain_idx > 0:
                            yield _sse_event("status", {"phase": "fallback_" + model_id})

                        gemini_client = genai.Client(api_key=_config.gemini_api_key)
                        api_model = MODEL_API_IDS.get(model_id, model_id)

                        gemini_contents = []
                        for msg in messages:
                            role = "user" if msg["role"] == "user" else "model"
                            content = msg["content"]
                            if isinstance(content, list):
                                text = " ".join(
                                    b["text"] for b in content
                                    if isinstance(b, dict) and b.get("type") == "text"
                                )
                                content = text if text else str(content)
                            gemini_contents.append(
                                genai_types.Content(
                                    role=role,
                                    parts=[genai_types.Part(text=content)],
                                )
                            )

                        # Google Search grounding (sin thinking — son incompatibles
                        # en Gemini 2.5 Pro: thinking impide que el modelo ejecute
                        # Google Search, resultando en grounding_metadata vacío).
                        gemini_tools = [
                            genai_types.Tool(google_search=genai_types.GoogleSearch())
                        ]

                        gemini_config = genai_types.GenerateContentConfig(
                            system_instruction=system_prompt,
                            max_output_tokens=32000,
                            temperature=0.7,
                            tools=gemini_tools,
                        )

                        yield _sse_event("status", {"phase": "thinking"})

                        # Non-streaming: grounding_metadata no se propaga en chunks
                        # de streaming con AFC (Automatic Function Calling) en google-genai SDK.
                        # Usamos generate_content y emitimos el texto como SSE.
                        gemini_response = await gemini_client.aio.models.generate_content(
                            model=api_model,
                            contents=gemini_contents,
                            config=gemini_config,
                        )

                        # Emitir texto como SSE (excluyendo thoughts)
                        if gemini_response.candidates and gemini_response.candidates[0].content:
                            for part in gemini_response.candidates[0].content.parts:
                                if part.text and not getattr(part, "thought", False):
                                    text_streamed = True
                                    yield _sse_event("text_delta", {"text": part.text})

                        # Extract usage
                        usage_meta = getattr(gemini_response, "usage_metadata", None)
                        stream_input_tokens = getattr(usage_meta, "prompt_token_count", 0) if usage_meta else 0
                        stream_output_tokens = getattr(usage_meta, "candidates_token_count", 0) if usage_meta else 0

                        # Extract web sources from grounding_metadata
                        # (model decides whether to use Google Search — sources are optional)
                        from pipeline.model_router import _extract_google_web_sources
                        gemini_web_sources = _extract_google_web_sources(gemini_response)

                        model_used = model_id
                        break  # success

                    elif provider == "openai":
                        # OpenAI: usa Responses API con web_search_preview
                        if chain_idx > 0:
                            yield _sse_event("status", {"phase": "fallback_" + model_id})

                        response = await _model_router.call_openai(
                            model_id, system_prompt, messages,
                            max_tokens=32000, web_search=True,
                        )

                        # Responses API: extraer texto y fuentes
                        from pipeline.model_router import (
                            _extract_openai_responses_text,
                            _extract_openai_web_sources,
                        )
                        oai_text = _extract_openai_responses_text(response)
                        oai_web_sources = _extract_openai_web_sources(response)
                        if oai_text:
                            text_streamed = True
                            yield _sse_event("text_delta", {"text": oai_text})

                        # Extract usage from Responses API
                        oai_usage = getattr(response, "usage", None)
                        stream_input_tokens = getattr(oai_usage, "input_tokens", 0) if oai_usage else 0
                        stream_output_tokens = getattr(oai_usage, "output_tokens", 0) if oai_usage else 0

                        model_used = model_id
                        break  # success

                except Exception as e:
                    is_overloaded = "overloaded" in str(e).lower()
                    logger.warning("Advisor stream: %s failed (overloaded=%s): %s",
                                   model_id, is_overloaded, str(e)[:200])
                    continue

            # 6. Fallback: if streaming produced nothing, extract from final
            if final is not None:
                logger.info("Advisor stream: text_streamed=%s, final blocks=%s",
                            text_streamed,
                            [getattr(b, "type", "?") for b in final.content])
                if not text_streamed:
                    for block in final.content:
                        if getattr(block, "type", None) == "text":
                            yield _sse_event("text_delta", {"text": block.text})
                            text_streamed = True
                            break

            if not text_streamed:
                logger.warning("Advisor stream: no text in response")
                yield _sse_event("text_delta", {
                    "text": "No se pudo generar una respuesta. Intenta reformular tu pregunta.",
                })

            # 7. Extract web sources from all providers
            web_sources: list[dict] = []
            # Anthropic: from final message
            if final is not None:
                for block in final.content:
                    if getattr(block, "type", None) == "web_search_tool_result":
                        for item in getattr(block, "content", []):
                            if getattr(item, "type", None) == "web_search_result":
                                web_sources.append({
                                    "title": getattr(item, "title", ""),
                                    "url": getattr(item, "url", ""),
                                })
            # Google: from grounding_metadata (set in google branch above)
            web_sources.extend(gemini_web_sources)
            # OpenAI: from Responses API annotations (set in openai branch above)
            web_sources.extend(oai_web_sources)

            seen_urls: set[str] = set()
            web_source_list = []
            for ws in web_sources:
                if ws["url"] and ws["url"] not in seen_urls:
                    seen_urls.add(ws["url"])
                    web_source_list.append({
                        "document_id": ws["url"],
                        "title": ws["title"],
                        "doc_type": "web",
                        "similarity": 0,
                        "scope": "web",
                        "excerpt": ws["url"],
                    })

            # 8. Record cost (tokens already set by Google/OpenAI branches above;
            #    for Anthropic, extract from final message)
            if final is not None:
                stream_input_tokens = final.usage.input_tokens
                stream_output_tokens = final.usage.output_tokens

            stream_duration = int((time.monotonic() - stream_start) * 1000)
            cost = calculate_cost(model_used, stream_input_tokens, stream_output_tokens)

            if _cost_guard:
                await _cost_guard.record(
                    model=model_used, service="advisor", operation="advisor_stream",
                    input_tokens=stream_input_tokens, output_tokens=stream_output_tokens,
                    duration_ms=stream_duration,
                    consultant_id=payload.consultant_id,
                    project_id=payload.project_id,
                    metadata={"web_searches": len(web_sources),
                              "tier": effective_tier,
                              "thinking_budget": thinking_budget,
                              "chain": chain},
                )

            logger.info(
                "Advisor stream: model=%s, tier=%s, RAG=%s, web=%s (%d), history=%d msgs, "
                "tokens=%d in + %d out, cost=$%.4f",
                model_used, effective_tier, has_rag_context, bool(web_sources), len(web_sources),
                len(payload.conversation_history),
                stream_input_tokens, stream_output_tokens, cost,
            )

            yield _sse_event("done", {
                "web_search_used": len(web_sources) > 0,
                "web_sources": web_source_list,
                "model_used": model_used,
                "cost_usd": cost,
                "input_tokens": stream_input_tokens,
                "output_tokens": stream_output_tokens,
            })

        except Exception as e:
            logger.error(f"Error in advisor stream: {e}")
            yield _sse_event("error", {"message": str(e)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ═══════════════════════════════════════════════════════════════
# ANALISIS MULTI-AGENTE - LangGraph
# ═══════════════════════════════════════════════════════════════

class AnalyzeRequest(BaseModel):
    project_id: str = Field(pattern=_UUID_PATTERN)
    agents: Optional[list[str]] = None  # None = all agents
    model_override: Optional[str] = None
    tier: Optional[str] = None
    consultant_id: Optional[str] = None

    @field_validator("tier")
    @classmethod
    def validate_tier(cls, v):
        if v is not None and v not in _VALID_TIERS:
            raise ValueError(f"tier debe ser 'standard' o 'pro_plus', recibido: {v!r}")
        return v

    @field_validator("model_override")
    @classmethod
    def validate_model(cls, v):
        if v is not None and v not in _VALID_MODELS:
            raise ValueError(f"model_override no reconocido: {v!r}")
        return v


@app.post("/api/analyze")
@limiter.limit("5/minute")
async def analyze_project(request: Request, payload: AnalyzeRequest):
    """
    Lanza el analisis multi-agente (LangGraph) para un proyecto.
    El consultor elige que agentes ejecutar via el campo 'agents'.
    Optimizador y Redactor siempre se ejecutan con los hallazgos disponibles.
    """
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    await _verify_project_ownership(payload.project_id, payload.consultant_id)

    from pipeline.agents import run_project_analysis

    try:
        result = await run_project_analysis(
            project_id=payload.project_id,
            supabase_url=_config.supabase_url,
            supabase_key=_config.supabase_service_key,
            anthropic_api_key=_config.anthropic_api_key,
            openai_api_key=_config.openai_api_key,
            gemini_api_key=_config.gemini_api_key,
            agents=payload.agents,
            model_override=payload.model_override or "",
            tier=payload.tier or "standard",
            consultant_id=payload.consultant_id or "",
        )
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en analisis del proyecto {payload.project_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# ANALISIS HITL - Sesiones + Plan + Execute + Round2
# ═══════════════════════════════════════════════════════════════

class SessionUpdate(BaseModel):
    phase: Optional[str] = None
    proposed_plan: Optional[dict] = None
    approved_plan: Optional[dict] = None
    consultant_instructions: Optional[str] = None
    agent_focus: Optional[dict] = None
    round1_results: Optional[dict] = None
    round2_results: Optional[dict] = None


@app.post("/api/analyze/session")
@limiter.limit("10/minute")
async def create_session(request: Request, project_id: str = Form(...), consultant_id: str = Form(...)):
    """Create a new HITL analysis session."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    if not re.match(_UUID_PATTERN, project_id):
        raise HTTPException(status_code=400, detail="project_id debe ser UUID valido")

    await _verify_project_ownership(project_id, consultant_id)

    from supabase._async.client import create_client as acreate_client
    sb = await acreate_client(_config.supabase_url, _config.supabase_service_key)

    result = await sb.table("analysis_sessions").insert({
        "project_id": project_id,
        "consultant_id": consultant_id,
        "phase": "planning",
    }).execute()

    if not result.data:
        raise HTTPException(status_code=500, detail="Error creating session")

    return result.data[0]


@app.get("/api/analyze/session/{project_id}")
@limiter.limit("30/minute")
async def get_session(request: Request, project_id: str, consultant_id: Optional[str] = None):
    """Get the latest active session for a project."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    if not re.match(_UUID_PATTERN, project_id):
        raise HTTPException(status_code=400, detail="project_id debe ser UUID valido")

    await _verify_project_ownership(project_id, consultant_id)

    from supabase._async.client import create_client as acreate_client
    sb = await acreate_client(_config.supabase_url, _config.supabase_service_key)

    result = await (
        sb.table("analysis_sessions")
        .select("*")
        .eq("project_id", project_id)
        .eq("consultant_id", consultant_id)
        .neq("phase", "complete")
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )

    if not result.data:
        return {"session": None}
    return {"session": result.data[0]}


@app.patch("/api/analyze/session/{session_id}")
async def update_session(session_id: str, request: SessionUpdate, consultant_id: Optional[str] = None):
    """Update a session's state (phase, plan, results, etc.)."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    if not re.match(_UUID_PATTERN, session_id):
        raise HTTPException(status_code=400, detail="session_id debe ser UUID valido")

    from supabase._async.client import create_client as acreate_client
    sb = await acreate_client(_config.supabase_url, _config.supabase_service_key)

    # Verificar que la sesion pertenece al consultor
    if consultant_id:
        session_check = await (
            sb.table("analysis_sessions")
            .select("id")
            .eq("id", session_id)
            .eq("consultant_id", consultant_id)
            .single()
            .execute()
        )
        if not session_check.data:
            raise HTTPException(status_code=403, detail="No tienes acceso a esta sesion")

    updates = {k: v for k, v in request.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")

    result = await (
        sb.table("analysis_sessions")
        .update(updates)
        .eq("id", session_id)
        .execute()
    )

    if not result.data:
        raise HTTPException(status_code=404, detail="Session not found")
    return result.data[0]


async def _cleanup_analysis_progress(project_id: str):
    """Delete progress rows for a completed analysis (they served their Realtime purpose)."""
    try:
        from supabase._async.client import create_client as acreate_client
        sb = await acreate_client(_config.supabase_url, _config.supabase_service_key)
        await sb.table("analysis_progress").delete().eq("project_id", project_id).execute()
    except Exception as e:
        logger.warning(f"Could not cleanup analysis_progress for {project_id}: {e}")


async def _verify_project_ownership(project_id: str, consultant_id: Optional[str]) -> None:
    """Verifica que el consultor es dueno del proyecto. Lanza 403 si no."""
    if not consultant_id:
        raise HTTPException(status_code=401, detail="consultant_id requerido")
    from supabase._async.client import create_client as acreate_client
    sb = await acreate_client(_config.supabase_url, _config.supabase_service_key)
    result = await (
        sb.table("projects")
        .select("id")
        .eq("id", project_id)
        .eq("consultant_id", consultant_id)
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")


class PlanRequest(BaseModel):
    project_id: str = Field(pattern=_UUID_PATTERN)
    consultant_id: Optional[str] = None


class ExecuteRequest(BaseModel):
    project_id: str = Field(pattern=_UUID_PATTERN)
    agents: list[str]
    consultant_instructions: str = Field(default="", max_length=10000)
    agent_focus: dict[str, str] = {}
    model_override: Optional[str] = None
    tier: Optional[str] = None
    consultant_id: Optional[str] = None

    @field_validator("agents")
    @classmethod
    def validate_agents(cls, v):
        invalid = [a for a in v if a not in _VALID_AGENTS]
        if invalid:
            raise ValueError(f"agentes no reconocidos: {invalid}")
        return v

    @field_validator("tier")
    @classmethod
    def validate_tier(cls, v):
        if v is not None and v not in _VALID_TIERS:
            raise ValueError(f"tier debe ser 'standard' o 'pro_plus', recibido: {v!r}")
        return v

    @field_validator("model_override")
    @classmethod
    def validate_model(cls, v):
        if v is not None and v not in _VALID_MODELS:
            raise ValueError(f"model_override no reconocido: {v!r}")
        return v


class Round2Request(BaseModel):
    project_id: str = Field(pattern=_UUID_PATTERN)
    agents: list[str]
    consultant_instructions: str = Field(default="", max_length=10000)
    agent_focus: dict[str, str] = {}
    previous_findings: list[dict] = Field(default=[], max_length=200)
    model_override: Optional[str] = None
    tier: Optional[str] = None
    consultant_id: Optional[str] = None

    @field_validator("agents")
    @classmethod
    def validate_agents(cls, v):
        invalid = [a for a in v if a not in _VALID_AGENTS]
        if invalid:
            raise ValueError(f"agentes no reconocidos: {invalid}")
        return v

    @field_validator("tier")
    @classmethod
    def validate_tier(cls, v):
        if v is not None and v not in _VALID_TIERS:
            raise ValueError(f"tier debe ser 'standard' o 'pro_plus', recibido: {v!r}")
        return v

    @field_validator("model_override")
    @classmethod
    def validate_model(cls, v):
        if v is not None and v not in _VALID_MODELS:
            raise ValueError(f"model_override no reconocido: {v!r}")
        return v


@app.post("/api/analyze/plan")
@limiter.limit("5/minute")
async def analyze_plan(request: Request, payload: PlanRequest):
    """Fase 0: Carga datos del proyecto y genera un plan de analisis inteligente."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    await _verify_project_ownership(payload.project_id, payload.consultant_id)

    from pipeline.agents import plan_analysis

    try:
        result = await plan_analysis(
            project_id=payload.project_id,
            supabase_url=_config.supabase_url,
            supabase_key=_config.supabase_service_key,
            anthropic_api_key=_config.anthropic_api_key,
        )
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error planificando analisis {payload.project_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/analyze/execute")
@limiter.limit("5/minute")
async def analyze_execute(request: Request, payload: ExecuteRequest):
    """Fase 2: Ejecuta el analisis con instrucciones del consultor."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    await _verify_project_ownership(payload.project_id, payload.consultant_id)

    from pipeline.agents import run_project_analysis

    try:
        result = await run_project_analysis(
            project_id=payload.project_id,
            supabase_url=_config.supabase_url,
            supabase_key=_config.supabase_service_key,
            anthropic_api_key=_config.anthropic_api_key,
            openai_api_key=_config.openai_api_key,
            gemini_api_key=_config.gemini_api_key,
            agents=payload.agents,
            consultant_instructions=payload.consultant_instructions,
            agent_focus=payload.agent_focus,
            model_override=payload.model_override or "",
            tier=payload.tier or "standard",
            consultant_id=payload.consultant_id or "",
        )
        # Cleanup old progress rows (Realtime already delivered them)
        await _cleanup_analysis_progress(payload.project_id)
        return result
    except HTTPException:
        raise
    except Exception as e:
        await _cleanup_analysis_progress(payload.project_id)
        logger.error(f"Error ejecutando analisis {payload.project_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/analyze/round2")
@limiter.limit("5/minute")
async def analyze_round2(request: Request, payload: Round2Request):
    """Fase 3: Segunda vuelta con hallazgos previos como contexto."""
    if _config is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    await _verify_project_ownership(payload.project_id, payload.consultant_id)

    from pipeline.agents import run_project_analysis

    try:
        result = await run_project_analysis(
            project_id=payload.project_id,
            supabase_url=_config.supabase_url,
            supabase_key=_config.supabase_service_key,
            anthropic_api_key=_config.anthropic_api_key,
            openai_api_key=_config.openai_api_key,
            gemini_api_key=_config.gemini_api_key,
            agents=payload.agents,
            consultant_instructions=payload.consultant_instructions,
            agent_focus=payload.agent_focus,
            round_number=2,
            previous_findings=payload.previous_findings,
            model_override=payload.model_override or "",
            tier=payload.tier or "standard",
            consultant_id=payload.consultant_id or "",
        )
        await _cleanup_analysis_progress(payload.project_id)
        return result
    except HTTPException:
        raise
    except Exception as e:
        await _cleanup_analysis_progress(payload.project_id)
        logger.error(f"Error en 2a vuelta {payload.project_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# KNOWLEDGE BASE - Gestión de documentos normativos generales
# ═══════════════════════════════════════════════════════════════

@app.get("/api/knowledge-base")
@limiter.limit("60/minute")
async def list_knowledge_base(
    request: Request,
    doc_type: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
):
    """Lista los documentos de la base de conocimiento general (normativa, BREFs, guías)."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()

    query = sb.table("knowledge_documents").select(
        "id, titulo, tipo, naturaleza_pdf, total_paginas, total_chunks, "
        "tablas_encontradas, metadata, estado, fecha_documento, fecha_ingesta"
    )

    if doc_type:
        query = query.eq("tipo", doc_type)

    if search:
        query = query.ilike("titulo", f"%{search}%")

    query = query.order("fecha_ingesta", desc=True)
    result = await query.execute()

    return {
        "documents": result.data or [],
        "total": len(result.data or []),
    }


@app.get("/api/knowledge-base/stats")
@limiter.limit("60/minute")
async def knowledge_base_stats(request: Request):
    """Estadísticas de la base de conocimiento general."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()

    # Total de documentos generales
    docs_result = await (
        sb.table("knowledge_documents")
        .select("id, tipo, total_chunks, total_paginas")
        .execute()
    )
    docs = docs_result.data or []

    # Contar por tipo
    by_type: dict[str, int] = {}
    total_chunks = 0
    total_pages = 0
    for doc in docs:
        tipo = doc.get("tipo", "desconocido")
        by_type[tipo] = by_type.get(tipo, 0) + 1
        total_chunks += doc.get("total_chunks") or 0
        total_pages += doc.get("total_paginas") or 0

    return {
        "total_documents": len(docs),
        "total_chunks": total_chunks,
        "total_pages": total_pages,
        "by_type": by_type,
    }


@app.delete("/api/knowledge-base/{doc_id}")
@limiter.limit("30/minute")
async def delete_knowledge_base_document(request: Request, doc_id: str):
    """Elimina un documento de la base de conocimiento general."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()

    # Verificar que el documento existe y es general
    doc_result = await (
        sb.table("knowledge_documents")
        .select("id, storage_path")
        .eq("id", doc_id)
        .execute()
    )

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    doc = doc_result.data[0]

    # Eliminar chunks (CASCADE debería hacerlo, pero por seguridad)
    await sb.table("knowledge_chunks").delete().eq("document_id", doc_id).execute()
    # Eliminar documento
    await sb.table("knowledge_documents").delete().eq("id", doc_id).execute()

    # Eliminar archivo de Storage si existe
    if doc.get("storage_path"):
        try:
            from pipeline import UnifiedIngestionService
            storage_svc = service.storage if service else None
            if storage_svc:
                await storage_svc.delete_file(doc["storage_path"])
        except Exception:
            pass  # No fallar si el archivo ya no existe

    return {"success": True, "deleted_id": doc_id}


# ═══════════════════════════════════════════════════════════════
# GOOGLE DRIVE - Conexión OAuth2 y gestión de carpetas
# ═══════════════════════════════════════════════════════════════

_gdrive_client_id = os.environ.get("GOOGLE_CLIENT_ID", "")
_gdrive_client_secret = os.environ.get("GOOGLE_CLIENT_SECRET", "")


def _gdrive_configured() -> bool:
    return bool(_gdrive_client_id and _gdrive_client_secret)


def _gdrive_redirect_uri() -> str:
    frontend = os.environ.get("FRONTEND_URL", "http://localhost:3000")
    return f"{frontend}/api/gdrive/callback"


@app.get("/api/gdrive/auth-url")
@limiter.limit("10/minute")
async def gdrive_auth_url(
    request: Request,
    consultant_id: str = Query(...),
    redirect_uri: Optional[str] = Query(None),
):
    """Generate Google OAuth2 authorization URL."""
    if not _gdrive_configured():
        raise HTTPException(
            status_code=501,
            detail="Google Drive no configurado. Falta GOOGLE_CLIENT_ID y GOOGLE_CLIENT_SECRET.",
        )

    from pipeline.google_drive import get_auth_url

    uri = redirect_uri or _gdrive_redirect_uri()
    url = get_auth_url(
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
        redirect_uri=uri,
        state=consultant_id,
    )
    return {"auth_url": url}


class GDriveExchangeRequest(BaseModel):
    code: str
    consultant_id: str
    redirect_uri: Optional[str] = None


@app.post("/api/gdrive/exchange")
@limiter.limit("10/minute")
async def gdrive_exchange(request: Request, payload: GDriveExchangeRequest):
    """
    Exchange OAuth code for tokens, save to DB, and create Drive folder structure.
    """
    if not _gdrive_configured():
        raise HTTPException(status_code=501, detail="Google Drive no configurado.")

    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    from pipeline.google_drive import exchange_code, GoogleDriveService

    # Exchange code for tokens (redirect_uri must match the one used in auth request)
    uri = payload.redirect_uri or _gdrive_redirect_uri()
    tokens = exchange_code(
        code=payload.code,
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
        redirect_uri=uri,
    )

    # Create folder structure in Drive
    gd = GoogleDriveService(
        access_token=tokens["access_token"],
        refresh_token=tokens["refresh_token"],
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
    )
    folders = gd.setup_full_structure()

    # Save tokens and folder IDs to database
    sb = await rag_service._get_supabase()
    await sb.table("consultant_gdrive").upsert({
        "consultant_id": payload.consultant_id,
        "access_token": tokens["access_token"],
        "refresh_token": tokens["refresh_token"],
        "token_expiry": tokens.get("token_expiry"),
        "root_folder_id": folders["root_folder_id"],
        "folder_mapping": folders,
    }).execute()

    return {
        "success": True,
        "root_folder_id": folders["root_folder_id"],
        "folders_created": len(folders),
    }


@app.get("/api/gdrive/picker-token")
@limiter.limit("20/minute")
async def gdrive_picker_token(request: Request, consultant_id: str = Query(...)):
    """Return a fresh access token for the Google Picker on the frontend."""
    if not _gdrive_configured():
        raise HTTPException(status_code=501, detail="Google Drive no configurado.")
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    result = await (
        sb.table("consultant_gdrive")
        .select("access_token, refresh_token")
        .eq("consultant_id", consultant_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Google Drive no conectado.")

    data = result.data[0]

    from pipeline.google_drive import GoogleDriveService

    gd = GoogleDriveService(
        access_token=data["access_token"],
        refresh_token=data["refresh_token"],
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
    )
    # The credentials may have been refreshed during construction
    fresh_token = gd.refreshed_token

    # Persist refreshed token if it changed
    if fresh_token != data["access_token"]:
        await (
            sb.table("consultant_gdrive")
            .update({"access_token": fresh_token})
            .eq("consultant_id", consultant_id)
            .execute()
        )

    return {"access_token": fresh_token, "client_id": _gdrive_client_id}


class GDriveSetupFoldersRequest(BaseModel):
    consultant_id: str
    root_folder_id: Optional[str] = None  # If provided, use as root (from Picker)


@app.post("/api/gdrive/setup-folders")
@limiter.limit("5/minute")
async def gdrive_setup_folders(request: Request, payload: GDriveSetupFoldersRequest):
    """
    Create Drive folder structure using tokens already saved in DB.
    If root_folder_id is provided (from Picker), use it as root.
    Runs in background and returns immediately. Poll /api/gdrive/status
    to check when root_folder_id is set.
    """
    if not _gdrive_configured():
        raise HTTPException(status_code=501, detail="Google Drive no configurado.")

    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    result = await (
        sb.table("consultant_gdrive")
        .select("access_token, refresh_token, root_folder_id")
        .eq("consultant_id", payload.consultant_id)
        .execute()
    )

    if not result.data:
        raise HTTPException(status_code=404, detail="Google Drive no conectado. Conecta primero.")

    data = result.data[0]

    # Skip if folder structure already exists (unless user is re-picking a folder)
    if data.get("root_folder_id") and not payload.root_folder_id:
        return {"status": "done", "root_folder_id": data["root_folder_id"], "already_exists": True}

    from pipeline.google_drive import GoogleDriveService

    gd = GoogleDriveService(
        access_token=data["access_token"],
        refresh_token=data["refresh_token"],
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
    )

    # If user picked a root folder via Picker, save it immediately
    if payload.root_folder_id:
        await (
            sb.table("consultant_gdrive")
            .update({"root_folder_id": payload.root_folder_id})
            .eq("consultant_id", payload.consultant_id)
            .execute()
        )

    # Fire-and-forget: run folder creation in background
    task = asyncio.create_task(
        _run_setup_folders(payload.consultant_id, gd, sb, payload.root_folder_id)
    )
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)

    return {
        "status": "running",
        "root_folder_id": payload.root_folder_id,
        "message": "Creando estructura de carpetas en segundo plano. Esto puede tardar 1-2 minutos.",
    }


async def _run_setup_folders(
    consultant_id: str,
    gd: "GoogleDriveService",  # noqa: F821
    sb: "AsyncClient",  # noqa: F821
    root_folder_id: Optional[str] = None,
) -> None:
    """Background task that creates the full folder structure in Google Drive."""
    try:
        folders = await asyncio.to_thread(gd.setup_full_structure, root_folder_id)

        await sb.table("consultant_gdrive").update({
            "root_folder_id": folders["root_folder_id"],
            "folder_mapping": folders,
        }).eq("consultant_id", consultant_id).execute()

        logger.info("Setup-folders complete for %s: root=%s", consultant_id, folders["root_folder_id"])
    except Exception as e:
        logger.error("Setup-folders failed for %s: %s", consultant_id, e)


@app.get("/api/gdrive/status")
@limiter.limit("30/minute")
async def gdrive_status(request: Request, consultant_id: str = Query(...)):
    """Check if the consultant has connected Google Drive."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    try:
        sb = await rag_service._get_supabase()
    except Exception as e:
        raise HTTPException(
            status_code=503,
            detail=f"Error conectando a Supabase: {e}. Verifica SUPABASE_SERVICE_ROLE_KEY.",
        )
    result = await (
        sb.table("consultant_gdrive")
        .select("root_folder_id, folder_mapping, created_at, updated_at")
        .eq("consultant_id", consultant_id)
        .execute()
    )

    if not result.data:
        return {"connected": False}

    data = result.data[0]
    return {
        "connected": True,
        "root_folder_id": data["root_folder_id"],
        "connected_at": data["created_at"],
        "configured": _gdrive_configured(),
    }


async def _get_gdrive_service(consultant_id: str):
    """Helper: load tokens from DB and return a GoogleDriveService instance."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")
    if not _gdrive_configured():
        raise HTTPException(status_code=501, detail="Google Drive no configurado.")

    sb = await rag_service._get_supabase()
    result = await (
        sb.table("consultant_gdrive")
        .select("access_token, refresh_token")
        .eq("consultant_id", consultant_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Google Drive no conectado.")

    from pipeline.google_drive import GoogleDriveService
    data = result.data[0]
    gd = GoogleDriveService(
        access_token=data["access_token"],
        refresh_token=data["refresh_token"],
        client_id=_gdrive_client_id,
        client_secret=_gdrive_client_secret,
    )
    return gd, sb


@app.get("/api/gdrive/browse")
@limiter.limit("30/minute")
async def gdrive_browse(
    request: Request,
    consultant_id: str = Query(...),
    folder_id: str = Query(...),
    page_token: Optional[str] = Query(None),
):
    """
    Browse a folder in the consultant's Google Drive.
    Returns items with sync status (whether each file is already indexed in the DB).
    """
    gd, sb = await _get_gdrive_service(consultant_id)

    listing = gd.list_folder(folder_id, page_token)

    # Get drive_file_ids already indexed in the DB
    file_ids = [item["id"] for item in listing["items"] if not item["isFolder"]]
    indexed_ids: set[str] = set()
    if file_ids:
        # Query in batches of 50 to avoid too-long queries
        for i in range(0, len(file_ids), 50):
            batch = file_ids[i : i + 50]
            indexed_result = await (
                sb.table("knowledge_documents")
                .select("drive_file_id")
                .in_("drive_file_id", batch)
                .execute()
            )
            for row in indexed_result.data or []:
                if row.get("drive_file_id"):
                    indexed_ids.add(row["drive_file_id"])

    # Enrich items with indexed status
    for item in listing["items"]:
        if item["isFolder"]:
            item["indexed"] = None
        else:
            item["indexed"] = item["id"] in indexed_ids

    return listing


class GDriveIngestRequest(BaseModel):
    consultant_id: str
    file_id: str
    file_name: str
    folder_path: str = ""  # breadcrumb path for context


@app.post("/api/gdrive/ingest-file")
@limiter.limit("10/minute")
async def gdrive_ingest_file(request: Request, payload: GDriveIngestRequest):
    """
    Download a file from Google Drive and ingest it through the pipeline.
    Stores drive_file_id on the resulting document for sync tracking.
    """
    if service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    gd, sb = await _get_gdrive_service(payload.consultant_id)

    # Check if already indexed
    existing = await (
        sb.table("knowledge_documents")
        .select("id")
        .eq("drive_file_id", payload.file_id)
        .execute()
    )
    if existing.data:
        raise HTTPException(
            status_code=409,
            detail="Este archivo ya esta indexado en la base de datos."
        )

    # Download from Drive
    try:
        file_bytes, filename, mime_type = gd.download_file(payload.file_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error descargando de Drive: {e}")

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="El archivo esta vacio.")

    if len(file_bytes) > 100 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Archivo demasiado grande (max 100 MB)")

    # Ingest through the pipeline
    try:
        result = await service.ingest(
            file_bytes=file_bytes,
            filename=filename,
            rag_scope="general",
        )

        if not result.success:
            raise HTTPException(
                status_code=422,
                detail=result.error or f"Error al procesar '{filename}'.",
            )

        # Update the document record with drive_file_id
        doc_id = result.supabase_doc_id or result.doc_id
        if doc_id:
            await (
                sb.table("knowledge_documents")
                .update({"drive_file_id": payload.file_id})
                .eq("id", doc_id)
                .execute()
            )

        return {
            **result.to_dict(),
            "drive_file_id": payload.file_id,
            "source_filename": filename,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# GOOGLE DRIVE SYNC - Auto-sync & bulk ingestion
# ═══════════════════════════════════════════════════════════════

class GDriveSyncRequest(BaseModel):
    consultant_id: str
    folder_id: Optional[str] = None  # None = use root folder


@app.post("/api/gdrive/sync")
@limiter.limit("3/minute")
async def gdrive_sync(request: Request, payload: GDriveSyncRequest):
    """
    Scan Google Drive for new documents and ingest them automatically.
    Creates a sync log entry, launches the heavy work as a background task,
    and returns immediately so the caller (Vercel) does not time out.
    The frontend polls /api/gdrive/sync-status to track progress.
    """
    if service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    try:
        gd, sb = await _get_gdrive_service(payload.consultant_id)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Sync: error getting GDrive service: %s", e)
        raise HTTPException(status_code=500, detail=f"Error obteniendo servicio GDrive: {e}")

    # Determine root folder
    folder_id = payload.folder_id
    if not folder_id:
        try:
            gdrive_row = await (
                sb.table("consultant_gdrive")
                .select("root_folder_id")
                .eq("consultant_id", payload.consultant_id)
                .execute()
            )
        except Exception as e:
            logger.error("Sync: error querying root folder: %s", e)
            raise HTTPException(status_code=500, detail=f"Error consultando carpeta raiz: {e}")
        if not gdrive_row.data or not gdrive_row.data[0].get("root_folder_id"):
            raise HTTPException(status_code=404, detail="No hay carpeta raiz configurada. Ve a Ajustes y pulsa 'Crear estructura de carpetas'.")
        folder_id = gdrive_row.data[0]["root_folder_id"]

    # Check if a sync is already running for this consultant
    try:
        running_check = await (
            sb.table("gdrive_sync_log")
            .select("id, started_at")
            .eq("consultant_id", payload.consultant_id)
            .eq("status", "running")
            .limit(1)
            .execute()
        )
    except Exception as e:
        logger.error("Sync: error checking running sync: %s", e)
        raise HTTPException(status_code=500, detail=f"Error consultando estado de sync: {e}")
    if running_check.data:
        # Auto-expire syncs that have been running for more than 120 minutes
        from datetime import datetime, timezone, timedelta
        started_at_str = running_check.data[0].get("started_at", "")
        stale_sync = False
        try:
            started_at = datetime.fromisoformat(started_at_str.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) - started_at > timedelta(minutes=120):
                stale_sync = True
        except (ValueError, TypeError):
            stale_sync = True  # Can't parse date — treat as stale

        if stale_sync:
            stale_id = running_check.data[0]["id"]
            logger.warning("Sync %s: stale sync detected (started %s), marking as error", stale_id, started_at_str)
            try:
                await (
                    sb.table("gdrive_sync_log")
                    .update({
                        "status": "error",
                        "completed_at": datetime.now(timezone.utc).isoformat(),
                        "error_message": "Sync expirado: superó el límite de 120 minutos. Posible caída del servidor.",
                    })
                    .eq("id", stale_id)
                    .execute()
                )
            except Exception:
                pass
            # Fall through to create a new sync
        else:
            return {
                "sync_id": running_check.data[0]["id"],
                "status": "already_running",
                "message": "Ya hay una sincronizacion en curso. Consulta sync-status para ver el progreso.",
            }

    # Create sync log entry
    try:
        sync_log = await (
            sb.table("gdrive_sync_log")
            .insert({
                "consultant_id": payload.consultant_id,
                "status": "running",
            })
            .execute()
        )
        if not sync_log.data:
            raise HTTPException(status_code=500, detail="No se pudo crear el registro de sync en gdrive_sync_log.")
        sync_id = sync_log.data[0]["id"]
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Sync: error creating sync log: %s", e)
        raise HTTPException(status_code=500, detail=f"Error creando registro de sync: {e}")

    # Fire-and-forget: launch heavy work in background
    async def _safe_sync_wrapper():
        """Wrapper to ensure exceptions from background task are always logged."""
        try:
            logger.info("Sync %s: background task STARTING", sync_id)
            await _run_sync_job(sync_id, payload.consultant_id, folder_id, gd, sb)
            logger.info("Sync %s: background task FINISHED", sync_id)
        except Exception as e:
            logger.error("Sync %s: background task CRASHED: %s", sync_id, e, exc_info=True)

    task = asyncio.create_task(_safe_sync_wrapper())
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)

    # Return immediately
    return {
        "sync_id": sync_id,
        "status": "running",
        "message": "Sincronizacion iniciada. Consulta sync-status para ver el progreso.",
    }


async def _run_sync_job(
    sync_id: str,
    consultant_id: str,
    folder_id: str,
    gd: "GoogleDriveService",  # noqa: F821
    sb: "AsyncClient",  # noqa: F821
) -> None:
    """Background task that does the actual sync work."""
    from datetime import datetime, timezone

    async def _update_sync_progress(**fields: object) -> None:
        """Helper to update gdrive_sync_log with current progress."""
        try:
            await (
                sb.table("gdrive_sync_log")
                .update(fields)
                .eq("id", sync_id)
                .execute()
            )
        except Exception:
            logger.warning("Sync %s: could not update progress", sync_id)

    try:
        # 1. Recursively list all files in Drive (run in thread to avoid blocking event loop)
        logger.info("Sync %s: _run_sync_job entered, folder=%s, service=%s", sync_id, folder_id, service is not None)
        logger.info("Sync %s: scanning Drive folder %s ...", sync_id, folder_id)
        all_files = await asyncio.to_thread(gd.list_all_files_recursive, folder_id)
        total_found = len(all_files)
        logger.info("Sync %s: found %d files in Drive", sync_id, total_found)

        # Update total_files_found immediately so the UI can show scan results
        await _update_sync_progress(total_files_found=total_found)

        # 2. Check which are already indexed (by drive_file_id OR titulo)
        all_drive_ids = [f["id"] for f in all_files]
        indexed_ids: set[str] = set()
        # 2a. Check by drive_file_id
        for i in range(0, len(all_drive_ids), 50):
            batch = all_drive_ids[i:i + 50]
            result = await (
                sb.table("knowledge_documents")
                .select("drive_file_id")
                .in_("drive_file_id", batch)
                .execute()
            )
            for row in result.data or []:
                if row.get("drive_file_id"):
                    indexed_ids.add(row["drive_file_id"])

        # 2b. Fallback: check by titulo (filename) for docs with NULL drive_file_id
        #     This catches docs ingested before drive_file_id was set
        all_filenames = [f["name"] for f in all_files]
        indexed_titles: set[str] = set()
        for i in range(0, len(all_filenames), 50):
            batch = all_filenames[i:i + 50]
            result = await (
                sb.table("knowledge_documents")
                .select("titulo, id, drive_file_id")
                .in_("titulo", batch)
                .execute()
            )
            for row in result.data or []:
                titulo = row.get("titulo")
                if titulo:
                    indexed_titles.add(titulo)
                    # Backfill drive_file_id if missing
                    if not row.get("drive_file_id"):
                        for f in all_files:
                            if f["name"] == titulo:
                                try:
                                    await (
                                        sb.table("knowledge_documents")
                                        .update({"drive_file_id": f["id"]})
                                        .eq("id", row["id"])
                                        .execute()
                                    )
                                    indexed_ids.add(f["id"])
                                    logger.info("Sync %s: backfilled drive_file_id for %s", sync_id, titulo)
                                except Exception:
                                    pass
                                break

        new_files = [f for f in all_files if f["id"] not in indexed_ids and f["name"] not in indexed_titles]

        # 2c. PDF priority: skip .md files when a .pdf version exists
        #     (in Drive, in indexed titles, or already in knowledge_documents)
        _pdf_basenames: set[str] = set()
        for f in all_files:
            if f["name"].lower().endswith(".pdf"):
                _pdf_basenames.add(f["name"][:-4])
        for title in indexed_titles:
            if title.lower().endswith(".pdf"):
                _pdf_basenames.add(title[:-4])
        try:
            _pdf_db = await (
                sb.table("knowledge_documents")
                .select("titulo")
                .like("titulo", "%.pdf")
                .execute()
            )
            for row in (_pdf_db.data or []):
                _pdf_basenames.add(row["titulo"][:-4])
        except Exception:
            pass

        _md_skipped_files = [
            f for f in new_files
            if f["name"].lower().endswith(".md")
            and f["name"][:-3] in _pdf_basenames
        ]
        new_files = [
            f for f in new_files
            if not (
                f["name"].lower().endswith(".md")
                and f["name"][:-3] in _pdf_basenames
            )
        ]
        if _md_skipped_files:
            logger.info("Sync %s: %d .md skipped (PDF version exists)", sync_id, len(_md_skipped_files))

        skipped = len(all_files) - len(new_files)
        logger.info("Sync %s: %d new files to ingest, %d already indexed", sync_id, len(new_files), skipped)

        # Update skipped count immediately
        await _update_sync_progress(files_skipped=skipped)

        # 3. Ingest new files (concurrent with semaphore)
        ingested = 0
        failed = 0
        details: list[dict] = [
            {"file": f["name"], "path": f.get("path", ""), "status": "skipped", "reason": "PDF version exists"}
            for f in _md_skipped_files
        ]
        _sync_lock = asyncio.Lock()
        _sem = asyncio.Semaphore(1)  # sequential to avoid memory crash (free(): invalid size)

        async def _ingest_one(file_info: dict) -> None:
            nonlocal ingested, failed, skipped
            fname = file_info["name"]
            fpath = file_info.get("path", "")
            try:
                file_bytes, filename, mime_type = await asyncio.to_thread(
                    gd.download_file, file_info["id"]
                )

                if len(file_bytes) == 0:
                    async with _sync_lock:
                        details.append({"file": fname, "status": "skipped", "reason": "empty file"})
                        skipped += 1
                        await _update_sync_progress(files_skipped=skipped)
                    return

                if len(file_bytes) > 100 * 1024 * 1024:
                    async with _sync_lock:
                        details.append({"file": fname, "status": "skipped", "reason": "too large (>100MB)"})
                        skipped += 1
                        await _update_sync_progress(files_skipped=skipped)
                    return

                result = await asyncio.wait_for(
                    service.ingest(file_bytes=file_bytes, filename=filename),
                    timeout=300,
                )

                if not result.success:
                    async with _sync_lock:
                        failed += 1
                        details.append({"file": fname, "path": fpath, "status": "error", "error": result.error or "Ingestion failed"})
                        logger.warning("Sync %s: ingestion failed for %s: %s", sync_id, fname, result.error)
                        await _update_sync_progress(files_failed=failed)
                    return

                doc_id = result.supabase_doc_id or result.doc_id
                if doc_id:
                    try:
                        await (
                            sb.table("knowledge_documents")
                            .update({"drive_file_id": file_info["id"]})
                            .eq("id", doc_id)
                            .execute()
                        )
                    except Exception as ue:
                        logger.warning("Sync %s: could not set drive_file_id for %s: %s", sync_id, fname, ue)

                # Auto-replace old .md version when a PDF is ingested
                if fname.lower().endswith(".pdf"):
                    _md_titulo = fname[:-4] + ".md"
                    try:
                        _old_md = await (
                            sb.table("knowledge_documents")
                            .select("id")
                            .eq("titulo", _md_titulo)
                            .execute()
                        )
                        for _md_row in (_old_md.data or []):
                            _md_id = _md_row["id"]
                            await sb.table("knowledge_chunks").delete().eq("document_id", _md_id).execute()
                            await sb.table("knowledge_documents").delete().eq("id", _md_id).execute()
                            logger.info("Sync %s: replaced .md '%s' with PDF '%s'", sync_id, _md_titulo, fname)
                            async with _sync_lock:
                                details.append({"file": _md_titulo, "status": "replaced", "reason": f"Replaced by PDF: {fname}"})
                    except Exception as _re:
                        logger.warning("Sync %s: could not cleanup old .md '%s': %s", sync_id, _md_titulo, _re)

                async with _sync_lock:
                    ingested += 1
                    details.append({"file": fname, "path": fpath, "status": "ingested", "document_id": doc_id, "chunks": result.num_chunks})
                    logger.info("Sync %s: ingested %s (%d chunks)", sync_id, fname, result.num_chunks or 0)
                    await _update_sync_progress(files_ingested=ingested)
                gc.collect()

            except (asyncio.TimeoutError, TimeoutError):
                async with _sync_lock:
                    failed += 1
                    details.append({"file": fname, "path": fpath, "status": "error", "error": "Timeout: >5 min"})
                    logger.warning("Sync %s: TIMEOUT processing %s (>5 min), skipping", sync_id, fname)
                    await _update_sync_progress(files_failed=failed)
                gc.collect()
            except Exception as e:
                async with _sync_lock:
                    failed += 1
                    details.append({"file": fname, "path": fpath, "status": "error", "error": str(e)[:200]})
                    logger.warning("Sync %s: failed %s: %s", sync_id, fname, e)
                    await _update_sync_progress(files_failed=failed)
                gc.collect()

        # Process files in concurrent batches
        for batch_start in range(0, len(new_files), 10):
            batch = new_files[batch_start:batch_start + 10]
            async def _bounded(fi: dict) -> None:
                async with _sem:
                    await _ingest_one(fi)
            await asyncio.gather(*[_bounded(fi) for fi in batch])

        # 4. Final update of sync log
        now_iso = datetime.now(timezone.utc).isoformat()
        await (
            sb.table("gdrive_sync_log")
            .update({
                "status": "completed",
                "completed_at": now_iso,
                "total_files_found": total_found,
                "files_ingested": ingested,
                "files_skipped": skipped,
                "files_failed": failed,
                "details": details,
            })
            .eq("id", sync_id)
            .execute()
        )

        # Update last_synced_at
        await (
            sb.table("consultant_gdrive")
            .update({"last_synced_at": now_iso})
            .eq("consultant_id", consultant_id)
            .execute()
        )

        logger.info(
            "Sync %s completed: %d ingested, %d skipped, %d failed",
            sync_id, ingested, skipped, failed,
        )

    except Exception as e:
        logger.error("Sync %s crashed: %s", sync_id, e)
        from datetime import datetime, timezone
        try:
            await (
                sb.table("gdrive_sync_log")
                .update({
                    "status": "error",
                    "completed_at": datetime.now(timezone.utc).isoformat(),
                    "error_message": str(e)[:500],
                })
                .eq("id", sync_id)
                .execute()
            )
        except Exception:
            logger.error("Sync %s: could not update error status", sync_id)


@app.get("/api/gdrive/sync-status")
@limiter.limit("30/minute")
async def gdrive_sync_status(request: Request, consultant_id: str = Query(...)):
    """
    Get sync status: last sync info + auto-sync setting.
    """
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()

    # Get GDrive config
    gdrive_result = await (
        sb.table("consultant_gdrive")
        .select("last_synced_at, auto_sync_enabled")
        .eq("consultant_id", consultant_id)
        .execute()
    )

    gdrive_config = gdrive_result.data[0] if gdrive_result.data else {}

    # Get last 5 sync logs
    logs_result = await (
        sb.table("gdrive_sync_log")
        .select("*")
        .eq("consultant_id", consultant_id)
        .order("started_at", desc=True)
        .limit(5)
        .execute()
    )

    # Check if a sync is currently running
    running = any(
        log.get("status") == "running" for log in (logs_result.data or [])
    )

    return {
        "last_synced_at": gdrive_config.get("last_synced_at"),
        "auto_sync_enabled": gdrive_config.get("auto_sync_enabled", True),
        "is_syncing": running,
        "recent_syncs": logs_result.data or [],
    }


class GDriveAutoSyncToggle(BaseModel):
    consultant_id: str
    enabled: bool


@app.post("/api/gdrive/sync-toggle")
@limiter.limit("10/minute")
async def gdrive_sync_toggle(request: Request, payload: GDriveAutoSyncToggle):
    """Toggle auto-sync on/off."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    await (
        sb.table("consultant_gdrive")
        .update({"auto_sync_enabled": payload.enabled})
        .eq("consultant_id", payload.consultant_id)
        .execute()
    )
    return {"success": True, "auto_sync_enabled": payload.enabled}


@app.delete("/api/gdrive/disconnect")
@limiter.limit("5/minute")
async def gdrive_disconnect(request: Request, consultant_id: str = Query(...)):
    """Disconnect Google Drive (delete stored tokens)."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    await (
        sb.table("consultant_gdrive")
        .delete()
        .eq("consultant_id", consultant_id)
        .execute()
    )

    return {"success": True}


# ─── RAG Health Check ──────────────────────────────────────────────
@app.get("/api/rag/health")
@limiter.limit("60/minute")
async def rag_health(request: Request):
    """Diagnóstico del sistema RAG: docs sin chunks, estadísticas."""
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    result = await sb.rpc("rag_health_check").execute()
    return {"status": "ok", "scopes": result.data or []}


# ─── Re-process documents ─────────────────────────────────────────
class ReprocessRequest(BaseModel):
    doc_ids: list[str]
    scope: str = "knowledge"  # "knowledge" or "project"

@app.post("/api/reprocess")
@limiter.limit("5/minute")
async def reprocess_documents(request: Request, req: ReprocessRequest):
    """Re-procesa documentos que no tienen chunks (o los tienen con error)."""
    if service is None or rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    sb = await rag_service._get_supabase()
    results = []

    for doc_id in req.doc_ids:
        try:
            # Obtener el documento original
            if req.scope == "knowledge":
                doc_result = await sb.table("knowledge_documents").select(
                    "id, titulo, storage_path, tipo"
                ).eq("id", doc_id).execute()
            else:
                doc_result = await sb.table("project_documents").select(
                    "id, titulo, storage_path, tipo, project_id"
                ).eq("id", doc_id).execute()

            if not doc_result.data:
                results.append({"doc_id": doc_id, "status": "not_found"})
                continue

            doc = doc_result.data[0]
            storage_path = doc.get("storage_path")

            if not storage_path:
                results.append({"doc_id": doc_id, "status": "no_storage_path"})
                continue

            # Descargar archivo desde Supabase Storage
            try:
                file_bytes = await sb.storage.from_("documentos").download(storage_path)
            except Exception as e:
                results.append({"doc_id": doc_id, "status": "download_error", "error": str(e)})
                continue

            filename = doc.get("titulo", "document.pdf")
            project_id = doc.get("project_id")
            rag_scope = "general" if req.scope == "knowledge" else "project"

            # Borrar chunks viejos antes de reprocesar
            chunk_table = "knowledge_chunks" if req.scope == "knowledge" else "project_chunks"
            await sb.table(chunk_table).delete().eq("document_id", doc_id).execute()

            # Re-ingestar usando la pipeline normal
            ingestion_result = await service.ingest(
                file_bytes=file_bytes,
                filename=filename,
                project_id=project_id,
                rag_scope=rag_scope,
            )

            results.append({
                "doc_id": doc_id,
                "status": "ok" if ingestion_result.success else "error",
                "num_chunks": ingestion_result.num_chunks,
                "error": ingestion_result.error,
            })
        except Exception as e:
            logger.exception(f"Error reprocesando {doc_id}: {e}")
            results.append({"doc_id": doc_id, "status": "error", "error": str(e)})

    success_count = sum(1 for r in results if r["status"] == "ok")
    return {
        "total": len(req.doc_ids),
        "success": success_count,
        "failed": len(req.doc_ids) - success_count,
        "results": results,
    }


# ─── Reclassify existing KB documents ────────────────────────────
@app.post("/api/knowledge-base/reclassify")
@limiter.limit("3/minute")
async def reclassify_knowledge_documents(request: Request):
    """
    Re-classify all knowledge_documents that have chunks,
    using the current KB classification logic (filename + content signals).
    Updates the 'tipo' field when the new classification differs.
    """
    if rag_service is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    from pipeline.classifier_chunker import KB_FILENAME_SIGNALS, KB_CONTENT_SIGNALS

    # DocType enum values → valid knowledge_documents.tipo values (CHECK constraint)
    DOCTYPE_TO_KB_TIPO = {
        "normativa": "legislacion",
        "ficha_seguridad": "documentacion_tecnica",
        "analisis_residuos": "clasificacion_residuos",
        "informe_certificacion": "referencia",
        "informe_tecnico": "documentacion_tecnica",
        "plan_gestion": "gestion_operativa",
    }

    sb = await rag_service._get_supabase()

    # 1. Get all knowledge_documents
    docs_result = await sb.table("knowledge_documents").select("id, titulo, tipo").execute()
    all_docs = docs_result.data or []
    if not all_docs:
        return {"total": 0, "reclassified": 0, "unchanged": 0, "changes": []}

    reclassified = 0
    unchanged = 0
    changes: list[dict] = []

    for doc in all_docs:
        doc_id = doc["id"]
        titulo = doc.get("titulo") or ""
        old_tipo = doc.get("tipo") or ""

        # Classify by filename
        filename_lower = titulo.lower().replace(" ", "_")
        new_tipo = None
        for doc_type_key, signals in KB_FILENAME_SIGNALS.items():
            if any(s in filename_lower for s in signals):
                raw = doc_type_key.value
                new_tipo = DOCTYPE_TO_KB_TIPO.get(raw, raw)
                break

        # If no filename match, classify by content from first 3 chunks
        if new_tipo is None:
            chunks_result = await (
                sb.table("knowledge_chunks")
                .select("contenido")
                .eq("document_id", doc_id)
                .order("chunk_index")
                .limit(3)
                .execute()
            )
            sample_text = " ".join(
                (c.get("contenido") or "").lower()
                for c in (chunks_result.data or [])
            )

            scores = {dt: 0 for dt in KB_CONTENT_SIGNALS}
            for dt, sigs in KB_CONTENT_SIGNALS.items():
                for sig in sigs:
                    if sig in sample_text:
                        scores[dt] += 1
            best = max(scores, key=scores.get)
            if scores[best] >= 2:
                raw = best.value
                new_tipo = DOCTYPE_TO_KB_TIPO.get(raw, raw)
            else:
                new_tipo = "legislacion"  # default KB

        if new_tipo != old_tipo:
            await (
                sb.table("knowledge_documents")
                .update({"tipo": new_tipo})
                .eq("id", doc_id)
                .execute()
            )
            changes.append({
                "doc_id": doc_id,
                "titulo": titulo,
                "old_tipo": old_tipo,
                "new_tipo": new_tipo,
            })
            reclassified += 1
            logger.info("Reclassified '%s': %s → %s", titulo, old_tipo, new_tipo)
        else:
            unchanged += 1

    return {
        "total": len(all_docs),
        "reclassified": reclassified,
        "unchanged": unchanged,
        "changes": changes,
    }


# ─── Advisor: Google Drive folder context (ephemeral) ─────────

DRIVE_CONTEXT_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv",
    ".txt", ".html", ".htm", ".md", ".json", ".xml",
}
DRIVE_MAX_FILES = 25
DRIVE_MAX_TOTAL_TEXT = 500_000  # ~500KB ≈ ~125K tokens (25 archivos × 20K)
DRIVE_MAX_PER_FILE = 20_000    # ~20KB per file


class DriveContextRequest(BaseModel):
    consultant_id: str
    folder_id: str
    recursive: bool = False


@app.post("/api/advisor/drive-context")
@limiter.limit("20/minute")
async def advisor_drive_context(request: Request, payload: DriveContextRequest):
    """
    Extract text from files in a Google Drive folder for ephemeral advisor context.
    Downloads, extracts text, and returns concatenated content — nothing is persisted.
    """
    gd, _sb = await _get_gdrive_service(payload.consultant_id)

    # 1. List files in folder
    if payload.recursive:
        all_files = gd.list_all_files_recursive(
            payload.folder_id,
            supported_extensions=DRIVE_CONTEXT_EXTENSIONS,
            max_folders=50,
        )
    else:
        listing = gd.list_folder(payload.folder_id)
        all_files = [
            f for f in listing["items"]
            if not f["isFolder"]
            and any(f["name"].lower().endswith(ext) for ext in DRIVE_CONTEXT_EXTENSIONS)
        ]

    if not all_files:
        return {
            "context_text": "",
            "files": [],
            "total_files_in_folder": 0,
            "processed_files": 0,
            "truncated": False,
        }

    # 2. Download and extract text (with limits)
    context_parts: list[str] = []
    file_metadata: list[dict] = []
    total_text = 0
    truncated = len(all_files) > DRIVE_MAX_FILES

    for f in all_files[:DRIVE_MAX_FILES]:
        if total_text >= DRIVE_MAX_TOTAL_TEXT:
            truncated = True
            break
        try:
            file_bytes, filename, _mime = await asyncio.to_thread(
                gd.download_file, f["id"]
            )
            if len(file_bytes) == 0:
                continue

            extracted = await asyncio.to_thread(
                _extract_binary_text, file_bytes, filename
            )
            if not extracted or extracted.startswith("["):
                continue

            budget = min(DRIVE_MAX_PER_FILE, DRIVE_MAX_TOTAL_TEXT - total_text)
            text_chunk = extracted[:budget]
            context_parts.append(f"=== {filename} ===\n{text_chunk}")
            total_text += len(text_chunk)
            file_metadata.append({
                "name": filename,
                "size": f.get("size") or len(file_bytes),
                "chars_extracted": len(text_chunk),
            })
        except Exception as e:
            logger.warning("drive-context: failed to process %s: %s", f.get("name", "?"), e)
            continue

    context_text = "\n\n".join(context_parts)

    return {
        "context_text": context_text,
        "files": file_metadata,
        "total_files_in_folder": len(all_files),
        "processed_files": len(file_metadata),
        "truncated": truncated,
    }


# ═══════════════════════════════════════════════════════════════
# COST MANAGEMENT ENDPOINTS
# ═══════════════════════════════════════════════════════════════


class CostLimitsUpdate(BaseModel):
    consultant_id: str
    anthropic_daily_limit: Optional[float] = None
    anthropic_monthly_limit: Optional[float] = None
    openai_daily_limit: Optional[float] = None
    openai_monthly_limit: Optional[float] = None
    google_daily_limit: Optional[float] = None
    google_monthly_limit: Optional[float] = None
    global_daily_limit: Optional[float] = None
    global_monthly_limit: Optional[float] = None
    alert_threshold_pct: Optional[int] = None
    auto_fallback: Optional[bool] = None
    block_on_global_limit: Optional[bool] = None


class ModelConfigUpdate(BaseModel):
    consultant_id: str
    service: str
    preferred_model: str
    fallback_chain: list[str] = []
    tier: str = "standard"


@app.get("/api/usage-stats")
@limiter.limit("60/minute")
async def usage_stats(
    request: Request,
    consultant_id: str = Query(...),
    days: int = Query(default=30),
):
    """Estadisticas de uso y costes para el dashboard."""
    if _cost_guard is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    stats = await _cost_guard.get_stats(consultant_id, days)
    return stats


@app.get("/api/cost-limits")
@limiter.limit("60/minute")
async def get_cost_limits(request: Request, consultant_id: str = Query(...)):
    """Obtener limites de coste del consultor."""
    if _cost_guard is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    from supabase import create_client
    sb = create_client(_config.supabase_url, _config.supabase_service_key)
    result = sb.table("consultant_cost_limits") \
        .select("*") \
        .eq("consultant_id", consultant_id) \
        .limit(1) \
        .execute()

    from pipeline.cost_guard import DEFAULT_LIMITS
    return result.data[0] if result.data else DEFAULT_LIMITS


@app.put("/api/cost-limits")
@limiter.limit("10/minute")
async def update_cost_limits(request: Request, payload: CostLimitsUpdate):
    """Actualizar limites de coste del consultor."""
    if _cost_guard is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    limits = {k: v for k, v in payload.model_dump().items()
              if k != "consultant_id" and v is not None}
    ok = await _cost_guard.update_limits(payload.consultant_id, limits)
    if not ok:
        raise HTTPException(status_code=500, detail="Error al actualizar limites")
    return {"status": "ok"}


@app.get("/api/model-config")
@limiter.limit("60/minute")
async def get_model_config(
    request: Request,
    consultant_id: str = Query(...),
    service: str = Query(default="advisor"),
):
    """Obtener config de modelo para un servicio."""
    if _cost_guard is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    config = await _cost_guard.get_model_config(consultant_id, service)
    if not config:
        defaults = SERVICE_DEFAULTS.get(service, {}).get("standard", {})
        return {
            "preferred_model": defaults.get("preferred_model", "claude-sonnet-4"),
            "fallback_chain": defaults.get("fallback_chain", []),
            "tier": "standard",
        }
    return config


@app.put("/api/model-config")
@limiter.limit("10/minute")
async def update_model_config(request: Request, payload: ModelConfigUpdate):
    """Actualizar config de modelo para un servicio."""
    if _cost_guard is None:
        raise HTTPException(status_code=503, detail="Service not initialized")

    ok = await _cost_guard.update_model_config(
        payload.consultant_id, payload.service,
        payload.preferred_model, payload.fallback_chain, payload.tier,
    )
    if not ok:
        raise HTTPException(status_code=500, detail="Error al actualizar config")
    return {"status": "ok"}


@app.get("/api/available-models")
@limiter.limit("60/minute")
async def available_models(request: Request):
    """Lista de modelos disponibles con precios y capacidades."""
    from pipeline.model_router import MODEL_CAPABILITIES
    models = []
    for model_id, pricing in MODEL_PRICING.items():
        if model_id == "text-embedding-3-large":
            continue  # No mostrar embeddings como opcion de LLM
        caps = MODEL_CAPABILITIES.get(model_id, {})
        models.append({
            "id": model_id,
            "provider": pricing["provider"],
            "input_price": pricing["input"],
            "output_price": pricing["output"],
            "thinking": caps.get("thinking", False),
            "web_search": caps.get("web_search", False),
            "vision": caps.get("vision", False),
            "max_tokens": caps.get("max_tokens", 8192),
            "context": caps.get("context", 200000),
        })
    return {"models": models, "defaults": SERVICE_DEFAULTS}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "api.server:app",
        host="0.0.0.0",
        port=int(os.environ.get("PORT", os.environ.get("API_PORT", "8000"))),
        reload=True,
    )
